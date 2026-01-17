import os
import io
import csv
import traceback
from datetime import datetime, timedelta, timezone
from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, make_response, send_file, Response
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, login_user, login_required, logout_user, current_user
import docx 
from functools import wraps
from sqlalchemy import or_, func, case
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Project internal imports
from config import Config
from models.database import db
from models.entities import User, Submission, Grade, LearningActivity, LearningGoal, Quiz, QuizDetail, Question, Course, Enrollment, PlatformSettings, AIIntegration, LMSIntegration
from services.ai_service import AIService
from services.ocr_service import OCRService
from services.grading_service import GradingService
from services.submission_service import SubmissionService
from services.quiz_service import QuizService
from services.notification_service import NotificationService
from services.activity_service import ActivityService
from services.feedback_service import FeedbackService
from services.goal_service import GoalService
from services.stats_service import StatsService
from services.report_service import ReportService
from repositories.quiz_repository import QuizRepository
from repositories.grade_repository import GradeRepository
from repositories.activity_repository import ActivityRepository
from repositories.goal_repository import GoalRepository
from repositories.feedback_repository import FeedbackRepository
from repositories.admin_repository import AdminRepository
from services.admin_service import AdminService

def create_app():
    app = Flask(__name__)
    app.config.from_object(Config)

    # Initialize Database
    db.init_app(app)
    
    # Helper function to convert UTC to GMT+3 (Turkey timezone)
    def utc_to_gmt3(utc_dt):
        """Convert UTC datetime to GMT+3 timezone"""
        if utc_dt is None:
            return None
        # GMT+3 is UTC+3
        gmt3_offset = timedelta(hours=3)
        if utc_dt.tzinfo is None:
            # If naive datetime, assume it's UTC
            utc_dt = utc_dt.replace(tzinfo=timezone.utc)
        return (utc_dt + gmt3_offset).replace(tzinfo=None)
    
    def get_gmt3_now():
        """Get current time in GMT+3"""
        return utc_to_gmt3(datetime.utcnow())

    # Login Manager Setup
    login_manager = LoginManager()
    login_manager.login_view = 'login' 
    login_manager.init_app(app)

    @login_manager.user_loader
    def load_user(user_id):
        return User.query.get(int(user_id))

    # --- GLOBAL USER INJECTION ---
    @app.context_processor
    def inject_user():
        return dict(user=current_user)
    
    # Add JSON filter for templates
    @app.template_filter('from_json')
    def from_json_filter(value):
        import json
        if not value:
            return {}
        try:
            return json.loads(value)
        except:
            return {}

    # Configure Upload Folders
    UPLOAD_FOLDER = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'static/uploads')
    app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
    os.makedirs(UPLOAD_FOLDER, exist_ok=True) 

    # Create Database Tables
    with app.app_context():
        # First, try to add missing columns (migration)
        try:
            import sqlite3
            # Use app.config instead of importing Config to avoid scope issues
            db_uri = app.config.get('SQLALCHEMY_DATABASE_URI', '')
            if not db_uri:
                raise ValueError("SQLALCHEMY_DATABASE_URI not found in app config")
            
            db_path = db_uri.replace('sqlite:///', '')
            
            if os.path.exists(db_path):
                conn = sqlite3.connect(db_path)
                cursor = conn.cursor()
                
                try:
                    # Check if learning_activity table exists
                    cursor.execute("""
                        SELECT name FROM sqlite_master 
                        WHERE type='table' AND name='learning_activity'
                    """)
                    table_exists = cursor.fetchone()
                    
                    if table_exists:
                        # Check if student_id column exists
                        cursor.execute("PRAGMA table_info(learning_activity)")
                        columns = [column[1] for column in cursor.fetchall()]
                        
                        if 'student_id' not in columns:
                            print("Adding student_id column to learning_activity table...")
                            cursor.execute("""
                                ALTER TABLE learning_activity 
                                ADD COLUMN student_id INTEGER 
                                REFERENCES users(id)
                            """)
                            conn.commit()
                            print("✓ Successfully added student_id column.")
                        else:
                            print("✓ student_id column already exists.")
                    else:
                        print("learning_activity table does not exist yet. Will be created by db.create_all()")
                    
                    # Check if quiz_details table exists and add explanation column if needed
                    cursor.execute("""
                        SELECT name FROM sqlite_master 
                        WHERE type='table' AND name='quiz_details'
                    """)
                    quiz_details_exists = cursor.fetchone()
                    
                    if quiz_details_exists:
                        cursor.execute("PRAGMA table_info(quiz_details)")
                        columns = [column[1] for column in cursor.fetchall()]
                        
                        if 'explanation' not in columns:
                            print("Adding explanation column to quiz_details table...")
                            cursor.execute("""
                                ALTER TABLE quiz_details 
                                ADD COLUMN explanation TEXT
                            """)
                            conn.commit()
                            print("✓ Successfully added explanation column.")
                        else:
                            print("✓ explanation column already exists.")
                    else:
                        print("quiz_details table does not exist yet. Will be created by db.create_all()")
                    
                    # Check if quizzes table exists and add category column if needed
                    cursor.execute("""
                        SELECT name FROM sqlite_master 
                        WHERE type='table' AND name='quizzes'
                    """)
                    quizzes_exists = cursor.fetchone()
                    
                    if quizzes_exists:
                        cursor.execute("PRAGMA table_info(quizzes)")
                        columns = [column[1] for column in cursor.fetchall()]
                        
                        if 'category' not in columns:
                            print("Adding category column to quizzes table...")
                            cursor.execute("""
                                ALTER TABLE quizzes 
                                ADD COLUMN category VARCHAR(50)
                            """)
                            conn.commit()
                            print("✓ Successfully added category column.")
                        else:
                            print("✓ category column already exists.")
                    else:
                        print("quizzes table does not exist yet. Will be created by db.create_all()")
                    
                    # Check if learning_goals table exists and migrate to new schema if needed
                    cursor.execute("""
                        SELECT name FROM sqlite_master 
                        WHERE type='table' AND name='learning_goals'
                    """)
                    goals_exists = cursor.fetchone()
                    
                    if goals_exists:
                        cursor.execute("PRAGMA table_info(learning_goals)")
                        columns = [column[1] for column in cursor.fetchall()]
                        
                        # Migrate from old schema (goal_name, target_value, current_value) to new schema (title, target_score, current_score, status)
                        if 'goal_name' in columns and 'title' not in columns:
                            print("Migrating learning_goals table to new schema...")
                            # Add new columns
                            if 'title' not in columns:
                                cursor.execute("ALTER TABLE learning_goals ADD COLUMN title VARCHAR(100)")
                            if 'target_score' not in columns:
                                cursor.execute("ALTER TABLE learning_goals ADD COLUMN target_score FLOAT")
                            if 'current_score' not in columns:
                                cursor.execute("ALTER TABLE learning_goals ADD COLUMN current_score FLOAT DEFAULT 0.0")
                            if 'status' not in columns:
                                cursor.execute("ALTER TABLE learning_goals ADD COLUMN status VARCHAR(20) DEFAULT 'In Progress'")
                            if 'updated_at' not in columns:
                                cursor.execute("ALTER TABLE learning_goals ADD COLUMN updated_at DATETIME")
                            if 'target_date' not in columns:
                                cursor.execute("ALTER TABLE learning_goals ADD COLUMN target_date DATETIME")
                            
                            # Copy data from old columns to new columns
                            cursor.execute("""
                                UPDATE learning_goals 
                                SET title = goal_name,
                                    target_score = CAST(target_value AS FLOAT),
                                    current_score = CAST(current_value AS FLOAT),
                                    status = 'In Progress',
                                    updated_at = created_at
                                WHERE title IS NULL OR target_score IS NULL
                            """)
                            conn.commit()
                            print("✓ Successfully migrated learning_goals table.")
                        else:
                            # Check if target_date column exists
                            if 'target_date' not in columns:
                                print("Adding target_date column to learning_goals table...")
                                cursor.execute("ALTER TABLE learning_goals ADD COLUMN target_date DATETIME")
                                conn.commit()
                                print("✓ Successfully added target_date column.")
                            else:
                                print("✓ learning_goals table already migrated or using new schema.")
                    else:
                        print("learning_goals table does not exist yet. Will be created by db.create_all()")
                except sqlite3.Error as e:
                    print(f"⚠ Migration warning: {e}")
                    conn.rollback()
                finally:
                    conn.close()
        except Exception as e:
            print(f"⚠ Migration check failed: {e}")
            print("   Continuing with db.create_all()...")
        
        # Create all tables (will update schema if needed)
        db.create_all()
        print("✓ Database tables created/updated successfully.")
        
        # Check for GEMINI_API_KEY
        from dotenv import load_dotenv
        load_dotenv()
        gemini_key = os.getenv('GEMINI_API_KEY')
        if not gemini_key:
            print("⚠ WARNING: GEMINI_API_KEY not found in environment variables.")
            print("   Please create a .env file with: GEMINI_API_KEY=your_key_here")
            print("   AI features will not work without this key.")
        else:
            print(f"✓ GEMINI_API_KEY loaded successfully (length: {len(gemini_key)})")
        
        # Auto-seed questions if database is empty
        from models.entities import Question
        if Question.query.count() == 0:
            try:
                from seed_questions import seed_questions
                seed_questions()
                print("✓ Questions automatically seeded on first run.")
            except Exception as e:
                print(f"⚠ Warning: Could not auto-seed questions: {e}")
                print("   Please run 'python seed_questions.py' manually to add questions.")

    # --- AUTHENTICATION CHECK & CACHE CONTROL ---
    @app.before_request
    def check_user_auth():
        public_routes = ['login', 'register', 'static', 'privacy', 'terms', 'index']
        # Export routes are protected by @login_required, so they will be handled correctly
        if not current_user.is_authenticated and request.endpoint not in public_routes:
            return redirect(url_for('login'))

    @app.after_request
    def add_header(response):
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        return response
    
    # Role Based Access Decorator
    def role_required(role):
        def wrapper(fn):
            @wraps(fn)
            @login_required
            def decorated_view(*args, **kwargs):
                if current_user.role != role:
                    flash(f"Access Denied: Only {role}s are authorized.", "danger")
                    return redirect(url_for('dashboard'))
                return fn(*args, **kwargs)
            return decorated_view
        return wrapper

    # --- AUTH ROUTES ---
    @app.route('/')
    def index():
        if current_user.is_authenticated:
            return redirect(url_for('dashboard'))
        return redirect(url_for('login'))

    @app.route('/register', methods=['POST'])
    def register():
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        role = request.form.get('role', 'Student') 
        
        # Username kontrolü
        if User.query.filter_by(username=username).first():
            flash("This username is already in use! Please choose a different username.", "danger")
            return redirect(url_for('login', mode='register'))
        
        # Email kontrolü
        if User.query.filter_by(email=email).first():
            flash("This email address is already registered! Please use a different email address.", "danger")
            return redirect(url_for('login', mode='register'))
        
        try:
            hashed_pw = generate_password_hash(password, method='pbkdf2:sha256')
            new_user = User(username=username, email=email, password=hashed_pw, role=role)
            db.session.add(new_user)
            db.session.commit()
            flash("Registration successful! You can now login.", "success")
            return redirect(url_for('login', registered=1))
        except Exception as e:
            db.session.rollback()
            flash("Registration failed. Please try again.", "danger")
            return redirect(url_for('login', mode='register'))

    @app.route('/login', methods=['GET', 'POST'])
    def login():
        # If already authenticated, redirect to appropriate dashboard
        if current_user.is_authenticated:
            if current_user.role == 'Admin':
                return redirect(url_for('admin_dashboard'))
            elif current_user.role == 'Instructor':
                return redirect(url_for('instructor_dashboard'))
            else:
                return redirect(url_for('dashboard'))
        
        # Clear flash messages on GET request (when coming from logout, but not from register)
        if request.method == 'GET' and not request.args.get('registered'):
            from flask import session
            session.pop('_flashes', None)
        
        if request.method == 'POST':
            email = request.form.get('email', '').strip()
            password = request.form.get('password', '')
            
            if not email or not password:
                flash("Please enter both email and password.", "danger")
                return render_template('login.html')
            
            user = User.query.filter_by(email=email).first()
            if user and check_password_hash(user.password, password):
                login_user(user)
                # Redirect based on role
                if user.role == 'Admin':
                    return redirect(url_for('admin_dashboard'))
                elif user.role == 'Instructor':
                    return redirect(url_for('instructor_dashboard'))
                else:
                    return redirect(url_for('dashboard'))
            flash("Invalid email or password.", "danger")
        return render_template('login.html')

    @app.route('/logout')
    @login_required
    def logout():
        # Clear all flash messages and session data before logout
        from flask import session
        session.clear()  # Clear all session data
        logout_user()
        return redirect(url_for('login'))

    @app.route('/dashboard')
    @login_required
    def dashboard():
        if current_user.role == 'Admin':
            return redirect(url_for('admin_dashboard'))
        if current_user.role == 'Instructor':
            return redirect(url_for('instructor_dashboard'))
        
        from datetime import timedelta
        
        # Get all submissions
        submissions = Submission.query.filter_by(student_id=current_user.id).order_by(Submission.created_at.asc()).all()
        
        # Calculate Speaking Score (average of pronunciation_score and fluency_score)
        speaking_subs = [s for s in submissions if s.submission_type == 'SPEAKING' and s.grade]
        speaking_score = 0.0
        if speaking_subs:
            scores = []
            for sub in speaking_subs:
                if sub.grade.pronunciation_score and sub.grade.fluency_score:
                    scores.append((sub.grade.pronunciation_score + sub.grade.fluency_score) / 2)
            speaking_score = round(sum(scores) / len(scores), 1) if scores else 0.0
        
        # Calculate Writing Score (average of writing submissions)
        writing_subs = [s for s in submissions if s.submission_type == 'WRITING' and s.grade]
        writing_score = round(sum(s.grade.score for s in writing_subs) / len(writing_subs), 1) if writing_subs else 0.0
        
        # Calculate Quiz Progress
        all_quizzes = Quiz.query.filter_by(user_id=current_user.id).all()
        completed_quizzes = len(all_quizzes)
        quiz_progress = completed_quizzes  # Can be enhanced with total available quizzes
        
        # Calculate Current Streak (consecutive days with submissions)
        current_streak = 0
        if submissions:
            # Get unique submission dates
            submission_dates = set()
            for sub in submissions:
                submission_dates.add(sub.created_at.date())
            
            # Calculate streak backwards from today
            today = datetime.utcnow().date()
            check_date = today
            while check_date in submission_dates:
                current_streak += 1
                check_date -= timedelta(days=1)
        
        # Calculate Weekly Goal Progress
        today = datetime.utcnow().date()
        week_start = today - timedelta(days=today.weekday())  # Monday of current week
        weekly_submissions = [s for s in submissions if s.created_at.date() >= week_start]
        weekly_goal_current = len(weekly_submissions)
        weekly_goal_target = 5  # Default weekly goal
        weekly_goal_percentage = min(100, int((weekly_goal_current / weekly_goal_target) * 100)) if weekly_goal_target > 0 else 0
        weekly_goal_remaining = max(0, weekly_goal_target - weekly_goal_current)
        
        # Get recent submissions for the chart
        recent_submissions = submissions[-10:] if len(submissions) > 10 else submissions
        
        # Calculate Handwritten Score
        handwritten_subs = [s for s in submissions if s.submission_type == 'HANDWRITTEN' and s.grade]
        
        # Prepare multi-line chart data: Speaking, Writing, Quiz, Handwritten scores by date
        from collections import defaultdict
        chart_data = {
            'dates': [],
            'speaking_scores': [],
            'writing_scores': [],
            'quiz_scores': [],
            'handwritten_scores': []
        }
        
        # Get Handwritten submissions
        handwritten_subs = [s for s in submissions if s.submission_type == 'HANDWRITTEN' and s.grade]
        
        # Collect all dates from submissions and quizzes
        all_dates = set()
        
        # Speaking submissions with dates
        for sub in speaking_subs:
            if sub.grade and sub.grade.pronunciation_score is not None and sub.grade.fluency_score is not None:
                date_key = sub.created_at.date()
                all_dates.add(date_key)
        
        # Writing submissions with dates
        for sub in writing_subs:
            if sub.grade and sub.grade.score is not None:
                date_key = sub.created_at.date()
                all_dates.add(date_key)
        
        # Handwritten submissions with dates
        for sub in handwritten_subs:
            if sub.grade and sub.grade.score is not None:
                date_key = sub.created_at.date()
                all_dates.add(date_key)
        
        # Quiz submissions with dates
        for quiz in all_quizzes:
            if quiz.date_taken and quiz.score is not None:
                date_key = quiz.date_taken.date() if isinstance(quiz.date_taken, datetime) else quiz.date_taken
                all_dates.add(date_key)
        
        # Sort dates
        sorted_dates = sorted(all_dates)
        
        # Create date-indexed dictionaries
        speaking_by_date = {}
        writing_by_date = {}
        handwritten_by_date = {}
        quiz_by_date = {}
        
        for sub in speaking_subs:
            if sub.grade and sub.grade.pronunciation_score is not None and sub.grade.fluency_score is not None:
                date_key = sub.created_at.date()
                score = (sub.grade.pronunciation_score + sub.grade.fluency_score) / 2
                if date_key not in speaking_by_date:
                    speaking_by_date[date_key] = []
                speaking_by_date[date_key].append(score)
        
        for sub in writing_subs:
            if sub.grade and sub.grade.score is not None:
                date_key = sub.created_at.date()
                if date_key not in writing_by_date:
                    writing_by_date[date_key] = []
                writing_by_date[date_key].append(sub.grade.score)
        
        for sub in handwritten_subs:
            if sub.grade and sub.grade.score is not None:
                date_key = sub.created_at.date()
                if date_key not in handwritten_by_date:
                    handwritten_by_date[date_key] = []
                handwritten_by_date[date_key].append(sub.grade.score)
        
        for quiz in all_quizzes:
            if quiz.date_taken and quiz.score is not None:
                date_key = quiz.date_taken.date() if isinstance(quiz.date_taken, datetime) else quiz.date_taken
                if date_key not in quiz_by_date:
                    quiz_by_date[date_key] = []
                quiz_by_date[date_key].append(quiz.score)
        
        # Average scores per date and build chart data
        for date in sorted_dates:
            chart_data['dates'].append(date.strftime('%d %b'))
            
            # Speaking: average if multiple submissions on same date
            if date in speaking_by_date:
                chart_data['speaking_scores'].append(round(sum(speaking_by_date[date]) / len(speaking_by_date[date]), 1))
            else:
                chart_data['speaking_scores'].append(0)  # Use 0 instead of None for better chart display
            
            # Writing: average if multiple submissions on same date
            if date in writing_by_date:
                chart_data['writing_scores'].append(round(sum(writing_by_date[date]) / len(writing_by_date[date]), 1))
            else:
                chart_data['writing_scores'].append(0)  # Use 0 instead of None
            
            # Handwritten: average if multiple submissions on same date
            if date in handwritten_by_date:
                chart_data['handwritten_scores'].append(round(sum(handwritten_by_date[date]) / len(handwritten_by_date[date]), 1))
            else:
                chart_data['handwritten_scores'].append(0)  # Use 0 instead of None
            
            # Quiz: average if multiple quizzes on same date
            if date in quiz_by_date:
                chart_data['quiz_scores'].append(round(sum(quiz_by_date[date]) / len(quiz_by_date[date]), 1))
            else:
                chart_data['quiz_scores'].append(0)  # Use 0 instead of None
        
        # Calculate Handwritten Score for insights
        handwritten_score = 0.0
        if handwritten_subs:
            handwritten_score = round(sum(s.grade.score for s in handwritten_subs) / len(handwritten_subs), 1)
        
        # Calculate Quiz Score for insights
        quiz_score = 0.0
        if all_quizzes:
            quiz_scores_list = [q.score for q in all_quizzes if q.score is not None]
            quiz_score = round(sum(quiz_scores_list) / len(quiz_scores_list), 1) if quiz_scores_list else 0.0
        
        # Determine AI Performance Insights (Strongest and Weakest areas)
        area_scores = {
            'Speaking': speaking_score,
            'Writing': writing_score,
            'Quiz': quiz_score,
            'Handwritten': handwritten_score
        }
        
        # Filter out zero scores for comparison
        non_zero_scores = {k: v for k, v in area_scores.items() if v > 0}
        
        if non_zero_scores:
            strongest_area = max(non_zero_scores, key=non_zero_scores.get)
            weakest_area = min(non_zero_scores, key=non_zero_scores.get)
            strongest_score = non_zero_scores[strongest_area]
            weakest_score = non_zero_scores[weakest_area]
        else:
            # If all scores are zero, show default values
            strongest_area = 'Speaking'
            weakest_area = 'Handwritten'
            strongest_score = 0.0
            weakest_score = 0.0
        
        # Determine Recommended Next Step
        recommended_next = "Start Your First Activity"
        recommended_link = "/assignments"
        if not speaking_subs:
            recommended_next = "Improve Your Speaking"
            recommended_link = "/speaking"
        elif not writing_subs:
            recommended_next = "Improve Your Writing"
            recommended_link = "/submit/writing"
        elif speaking_score < 70:
            recommended_next = "Improve Your Speaking"
            recommended_link = "/speaking"
        elif writing_score < 70:
            recommended_next = "Improve Your Writing"
            recommended_link = "/submit/writing"
        elif completed_quizzes == 0:
            recommended_next = "Take a Quiz"
            recommended_link = "/quizzes"
        
        # Get latest graded submission for recommendations
        latest_graded = Submission.query.filter_by(student_id=current_user.id).join(Grade).order_by(Submission.created_at.desc()).first()
        
        # Get recommendations using StatsService
        recommendations = StatsService.fetch_recommendations(current_user.id)
        
        # Get adaptive insights (UC17)
        from services.adaptive_insights_service import AdaptiveInsightsService
        adaptive_insights = AdaptiveInsightsService.get_active_insights(current_user.id)
        
        # Get user goals using GoalService
        user_goals = GoalService.get_user_goals(current_user.id)[:2]
        
        # Calculate pending tasks - activities assigned to this student
        if current_user.role == 'Student':
            from services.activity_service import ActivityService
            student_activities = ActivityService.get_activities_for_student(current_user.id)
            # Get submitted activity IDs
            submitted_activity_ids = set(s.activity_id for s in submissions if s.activity_id)
            # Count activities not yet submitted
            pending_activities = [a for a in student_activities if a.id not in submitted_activity_ids]
            pending_count = len(pending_activities)
            
            # Get upcoming deadlines - activities with due_date in the future
            now = datetime.utcnow()
            upcoming_deadlines = [
                a for a in student_activities 
                if a.due_date and a.due_date >= now and a.id not in submitted_activity_ids
            ]
            # Sort by due_date (earliest first)
            upcoming_deadlines.sort(key=lambda x: x.due_date)
            # Limit to 5 most upcoming
            upcoming_deadlines = upcoming_deadlines[:5]
        else:
            # For instructors/admins, count all upcoming activities (for class performance monitoring - FR14)
            # Filter activities with due dates in the future or no due date (ongoing activities)
            pending_activities = LearningActivity.query.filter(
                or_(
                    LearningActivity.due_date >= datetime.utcnow(),
                    LearningActivity.due_date == None
                )
            ).order_by(LearningActivity.due_date.asc()).all()
            pending_count = len(pending_activities)
            upcoming_deadlines = []
        
        # Calculate total submissions
        total_submissions = len(submissions)
        
        # Calculate average score across all graded submissions
        graded_subs = [s for s in submissions if s.grade]
        avg_score = round(sum(s.grade.score for s in graded_subs) / len(graded_subs), 1) if graded_subs else 0.0
        
        return render_template('dashboard.html', 
                               submissions=submissions,
                               recent_submissions=recent_submissions,
                               speaking_score=speaking_score,
                               writing_score=writing_score,
                               quiz_progress=quiz_progress,
                               current_streak=current_streak,
                               weekly_goal_current=weekly_goal_current,
                               weekly_goal_target=weekly_goal_target,
                               weekly_goal_percentage=weekly_goal_percentage,
                               weekly_goal_remaining=weekly_goal_remaining,
                               recommended_next=recommended_next,
                               recommended_link=recommended_link,
                               latest_graded=latest_graded,
                               goals=user_goals,
                               speaking_subs=speaking_subs,
                               writing_subs=writing_subs,
                               has_chart_data=len(recent_submissions) > 0,
                               chart_data=chart_data,
                               pending_count=pending_count,
                               total_submissions=total_submissions,
                               average_score=avg_score,
                               strongest_area=strongest_area,
                               strongest_score=strongest_score,
                               weakest_area=weakest_area,
                               weakest_score=weakest_score,
                               recommendations=recommendations,
                               adaptive_insights=adaptive_insights,
                               upcoming_deadlines=upcoming_deadlines)

    @app.route('/courses')
    @login_required
    def student_courses():
        from models.entities import Enrollment, Course, LearningActivity, Submission, Grade
        # Get courses where this student is enrolled
        enrollments = Enrollment.query.filter_by(student_id=current_user.id, status='active').all()
        enrolled_courses = []
        
        # Get all student submissions for this student
        user_subs = Submission.query.filter_by(student_id=current_user.id).all()
        submitted_activity_ids = set(s.activity_id for s in user_subs if s.activity_id)
        submissions_with_grades = {s.activity_id: s for s in user_subs if s.activity_id and s.grade}
        
        for enrollment in enrollments:
            course = Course.query.get(enrollment.course_id)
            if course and course.is_active:
                # Get assignments for this course
                course_assignments = LearningActivity.query.filter(
                    LearningActivity.courses.any(id=course.id)
                ).all()
                
                # Calculate statistics
                total_assignments = len(course_assignments)
                
                # Completed: assignments with submitted and graded submissions
                completed_count = 0
                pending_count = 0
                scores = []
                
                for assignment in course_assignments:
                    if assignment.id in submitted_activity_ids:
                        submission = submissions_with_grades.get(assignment.id)
                        if submission and submission.grade and submission.grade.instructor_approved:
                            completed_count += 1
                            scores.append(submission.grade.score)
                        else:
                            pending_count += 1
                    else:
                        pending_count += 1
                
                # Calculate average score
                avg_score = round(sum(scores) / len(scores), 1) if scores else 0.0
                
                enrolled_courses.append({
                    'course': course,
                    'enrolled_at': enrollment.enrolled_at,
                    'total_assignments': total_assignments,
                    'completed_count': completed_count,
                    'pending_count': pending_count,
                    'avg_score': avg_score
                })
        
        return render_template('student_courses.html', enrolled_courses=enrolled_courses)

    @app.route('/courses/<int:course_id>')
    @login_required
    def student_course_detail(course_id):
        from models.entities import Course, Enrollment, LearningActivity, Submission
        # Get course
        course = Course.query.get_or_404(course_id)
        if not course.is_active:
            flash('Course is not active.', 'danger')
            return redirect(url_for('student_courses'))
        
        # Verify student is enrolled in this course
        enrollment = Enrollment.query.filter_by(
            student_id=current_user.id, 
            course_id=course_id, 
            status='active'
        ).first()
        
        if not enrollment:
            flash('You are not enrolled in this course.', 'danger')
            return redirect(url_for('student_courses'))
        
        # Get assignments for this course
        course_assignments = LearningActivity.query.filter(
            LearningActivity.courses.any(id=course_id)
        ).order_by(LearningActivity.due_date.asc(), LearningActivity.created_at.desc()).all()
        
        # Get student submissions to check completion status
        user_subs = Submission.query.filter_by(student_id=current_user.id).all()
        submitted_ids = set(s.activity_id for s in user_subs if s.activity_id and s.activity_id is not None)
        submissions_with_grades = {s.activity_id: s for s in user_subs if s.activity_id and s.grade}
        
        # Prepare assignment data with status
        assignments_data = []
        now = datetime.utcnow()
        for assignment in course_assignments:
            status = 'pending'
            if assignment.id in submitted_ids:
                if assignment.id in submissions_with_grades:
                    status = 'graded'
                else:
                    status = 'submitted'
            
            # Check if overdue
            is_overdue = False
            if assignment.due_date and assignment.due_date < now and status == 'pending':
                is_overdue = True
            
            assignments_data.append({
                'assignment': assignment,
                'status': status,
                'is_overdue': is_overdue,
                'submission': submissions_with_grades.get(assignment.id)
            })
        
        return render_template('student_course_detail.html', 
                             course=course, 
                             enrolled_at=enrollment.enrolled_at,
                             assignments_data=assignments_data)

    @app.route('/assignments/<int:activity_id>/view')
    @login_required
    def student_assignment_detail(activity_id):
        """Student view of assignment details before starting"""
        from models.entities import LearningActivity, Question, Enrollment
        from sqlalchemy import func
        
        # Get assignment
        assignment = LearningActivity.query.get_or_404(activity_id)
        
        # Verify student has access (enrolled in at least one course with this assignment)
        if current_user.role == 'Student':
            # Check if student is enrolled in any course that has this assignment
            student_enrollments = Enrollment.query.filter_by(
                student_id=current_user.id,
                status='active'
            ).all()
            enrolled_course_ids = [e.course_id for e in student_enrollments]
            
            assignment_course_ids = [c.id for c in assignment.courses] if assignment.courses else []
            
            # Assignment must have at least one course
            if not assignment_course_ids:
                flash('This assignment is not assigned to any course.', 'danger')
                return redirect(url_for('view_assignments'))
            
            # Check if there's any overlap between student's enrolled courses and assignment's courses
            has_access = False
            if assignment.student_id is None:  # Assigned to all students in the courses
                has_access = any(cid in enrolled_course_ids for cid in assignment_course_ids)
            else:  # Assigned to specific student
                has_access = (assignment.student_id == current_user.id and 
                             any(cid in enrolled_course_ids for cid in assignment_course_ids))
            
            if not has_access:
                flash('You do not have access to this assignment.', 'danger')
                return redirect(url_for('view_assignments'))
        
        # Prepare type-specific data
        writing_prompt = None
        word_limit = None
        speaking_prompt = None
        min_duration = None
        reference_image = None
        passage_text = None
        question_count = None
        
        # For quiz, get question count
        if assignment.activity_type == 'QUIZ' and assignment.quiz_category:
            question_count = Question.query.filter(
                func.lower(Question.category) == func.lower(assignment.quiz_category)
            ).count()
        
        # Note: writing_prompt, word_limit, speaking_prompt, min_duration, reference_image, passage_text
        # are not currently stored in the database. They would need to be added to LearningActivity model.
        # For now, we'll show what's available (description, quiz_category, etc.)
        
        return render_template('student_assignment_detail.html',
                             assignment=assignment,
                             writing_prompt=writing_prompt,
                             word_limit=word_limit,
                             speaking_prompt=speaking_prompt,
                             min_duration=min_duration,
                             reference_image=reference_image,
                             passage_text=passage_text,
                             question_count=question_count)

    @app.route('/assignments')
    @login_required
    def view_assignments():
        now = datetime.utcnow()

        # Filter assignments for students based on student_id
        if current_user.role == 'Student':
            # Get activities assigned to this student (student_id is None for all students, or matches current_user.id)
            from services.activity_service import ActivityService
            all_activities = ActivityService.get_activities_for_student(current_user.id)
        else:
            # For instructors/admins, show all activities
            all_activities = LearningActivity.query.order_by(LearningActivity.due_date.asc()).all()

        # Student submissions to mark completed assignments (including quiz submissions)
        user_subs = Submission.query.filter_by(student_id=current_user.id).all()
        submitted_ids = set(s.activity_id for s in user_subs if s.activity_id and s.activity_id is not None)
        
        # Get submissions with their grades for status determination
        submissions_with_grades = {s.activity_id: s for s in user_subs if s.activity_id and s.grade}

        # Filter assignments for students
        if current_user.role == 'Student':
            # Get all active assignments (not expired)
            active_activities = [a for a in all_activities if not a.due_date or a.due_date >= now]
            
            # Categorize assignments:
            # Active: not submitted yet
            active_activities_list = [a for a in active_activities if a.id not in submitted_ids]
            
            # Pending: submitted but waiting for instructor approval (has grade but not approved)
            pending_activities = []
            completed_activities = []
            
            for activity_id in submitted_ids:
                activity = next((a for a in all_activities if a.id == activity_id), None)
                if activity:
                    submission = submissions_with_grades.get(activity_id)
                    if submission and submission.grade:
                        if submission.grade.instructor_approved:
                            completed_activities.append(activity)
                        else:
                            pending_activities.append(activity)
                    else:
                        # Submitted but no grade yet - treat as pending
                        pending_activities.append(activity)
            
            # Default: show active
            activities = active_activities_list
        else:
            activities = all_activities
            active_activities_list = activities
            pending_activities = []
            completed_activities = []

        # Calculate counts for display
        active_count = len(active_activities_list)
        pending_count = len(pending_activities)
        completed_count = len(completed_activities)

        return render_template('assignments.html', 
                               activities=activities,
                               active_activities=active_activities_list,
                               pending_activities=pending_activities,
                               completed_activities=completed_activities,
                               active_count=active_count,
                               pending_count=pending_count,
                               completed_count=completed_count,
                               now=now,
                               submitted_ids=submitted_ids,
                               user_subs=user_subs,
                               submissions_with_grades=submissions_with_grades)

    @app.route('/instructor/assignments')
    @role_required('Instructor')
    def instructor_assignments():
        now = datetime.utcnow()
        
        # Use SQL with LEFT JOIN + GROUP BY for efficient stats calculation
        # Calculate: Total submissions, Graded (instructor_approved=True), Pending (instructor_approved=False or NULL)
        stats_query = db.session.query(
            LearningActivity,
            func.count(func.distinct(Submission.id)).label('total_submissions'),
            func.sum(case((Grade.instructor_approved == True, 1), else_=0)).label('graded_submissions'),
            func.sum(case((Grade.instructor_approved == False, 1), else_=0)).label('pending_submissions')
        ).outerjoin(Submission, Submission.activity_id == LearningActivity.id)\
         .outerjoin(Grade, Grade.submission_id == Submission.id)\
         .group_by(LearningActivity.id)\
         .order_by(LearningActivity.due_date.asc()).all()
        
        # Build activity_stats list from query results
        activity_stats = []
        for activity, total, graded, pending in stats_query:
            # Handle None values from SQL (when no submissions exist)
            total_submissions = total or 0
            graded_submissions = int(graded) if graded else 0
            # Pending = submissions with grade but instructor_approved = False
            pending_submissions = int(pending) if pending else 0
            
            # Get courses for this activity
            # activity.courses is already a list (InstrumentedList), no need for .all()
            activity_courses = list(activity.courses) if hasattr(activity, 'courses') and activity.courses else []
            
            activity_stats.append({
                'activity': activity,
                'total_submissions': total_submissions,
                'graded_submissions': graded_submissions,
                'pending_submissions': pending_submissions,
                'courses': activity_courses
            })
        
        return render_template('instructor_assignments.html', 
                             activity_stats=activity_stats, 
                             now=now)

    @app.route('/instructor/assignments/create', methods=['GET', 'POST'])
    @role_required('Instructor')
    def instructor_create_assignment():
        # Get students enrolled in courses taught by this instructor
        from models.entities import User, Course, Enrollment
        # Get courses where this instructor teaches
        instructor_courses = Course.query.filter_by(instructor_id=current_user.id, is_active=True).all()
        course_ids = [c.id for c in instructor_courses]
        
        # Get all enrollments for these courses
        enrolled_students = []
        if course_ids:
            enrollments = Enrollment.query.filter(
                Enrollment.course_id.in_(course_ids),
                Enrollment.status == 'active'
            ).all()
            # Get unique student IDs
            student_ids = list(set([e.student_id for e in enrollments]))
            # Get student objects
            enrolled_students = User.query.filter(
                User.id.in_(student_ids),
                User.role == 'Student'
            ).order_by(User.username.asc()).all()
        
        # Use enrolled students, fallback to all students if no courses assigned
        all_students = enrolled_students if enrolled_students else User.query.filter_by(role='Student').order_by(User.username.asc()).all()
        
        # Build course-student mapping for JavaScript
        course_student_map = {}
        for course in instructor_courses:
            enrollments = Enrollment.query.filter_by(
                course_id=course.id,
                status='active'
            ).all()
            student_ids = [e.student_id for e in enrollments]
            students = User.query.filter(
                User.id.in_(student_ids),
                User.role == 'Student'
            ).order_by(User.username.asc()).all()
            course_student_map[course.id] = [
                {'id': s.id, 'username': s.username, 'email': s.email}
                for s in students
            ]
        
        if request.method == 'POST':
            title = request.form.get('title')
            activity_type = request.form.get('activity_type')
            due_date_str = request.form.get('due_date')
            description = request.form.get('description')
            quiz_category = request.form.get('quiz_category') if activity_type == 'QUIZ' else None
            student_id_str = request.form.get('student_id', '').strip()
            
            # Get course IDs from form (multiple selection)
            course_ids_list = request.form.getlist('course_ids')
            course_ids_list = [int(cid) for cid in course_ids_list if cid and cid.strip()]
            
            # Validate that at least one course is selected
            if not course_ids_list:
                flash('Please select at least one course.', 'danger')
                return render_template('instructor_assignment_create.html', 
                                     students=all_students, 
                                     courses=instructor_courses,
                                     course_student_map=course_student_map)
            
            # Verify all selected courses belong to this instructor
            valid_courses = Course.query.filter(
                Course.id.in_(course_ids_list),
                Course.instructor_id == current_user.id,
                Course.is_active == True
            ).all()
            if len(valid_courses) != len(course_ids_list):
                flash('Invalid course selection.', 'danger')
                return render_template('instructor_assignment_create.html', 
                                     students=all_students, 
                                     courses=instructor_courses,
                                     course_student_map=course_student_map)
            
            # Parse student_id - empty string means assign to all students
            student_id = None
            if student_id_str and student_id_str != '':
                try:
                    student_id = int(student_id_str)
                    # Verify student exists and is enrolled in selected courses
                    student = User.query.get(student_id)
                    if not student or student.role != 'Student':
                        flash('Invalid student selected.', 'danger')
                        return render_template('instructor_assignment_create.html', 
                                             students=all_students, 
                                             courses=instructor_courses,
                                             course_student_map=course_student_map)
                    
                    # Verify student is enrolled in at least one selected course
                    student_enrollments = Enrollment.query.filter(
                        Enrollment.student_id == student_id,
                        Enrollment.course_id.in_(course_ids_list),
                        Enrollment.status == 'active'
                    ).first()
                    if not student_enrollments:
                        flash('Selected student is not enrolled in any of the selected courses.', 'danger')
                        return render_template('instructor_assignment_create.html', 
                                             students=all_students, 
                                             courses=instructor_courses,
                                             course_student_map=course_student_map)
                except ValueError:
                    flash('Invalid student ID format.', 'danger')
                    return render_template('instructor_assignment_create.html', 
                                         students=all_students, 
                                         courses=instructor_courses,
                                         course_student_map=course_student_map)

            if not title or not activity_type:
                flash('Title and type are required.', 'danger')
                return render_template('instructor_assignment_create.html', 
                                     students=all_students, 
                                     courses=instructor_courses,
                                     course_student_map=course_student_map)

            due_date = None
            if due_date_str:
                try:
                    due_date = datetime.strptime(due_date_str, '%Y-%m-%d')
                except ValueError:
                    flash('Invalid date format. Use YYYY-MM-DD.', 'danger')
                    return render_template('instructor_assignment_create.html', 
                                         students=all_students, 
                                         courses=instructor_courses,
                                         course_student_map=course_student_map)

            # Handle attachment file upload
            attachment_path = None
            attachment_filename = None
            if 'attachment' in request.files:
                attachment_file = request.files['attachment']
                if attachment_file and attachment_file.filename:
                    # Validate file size (max 10MB)
                    if attachment_file.content_length and attachment_file.content_length > 10 * 1024 * 1024:
                        flash('Attachment file size must be less than 10MB.', 'danger')
                        return render_template('instructor_assignment_create.html', 
                                             students=all_students, 
                                             courses=instructor_courses,
                                             course_student_map=course_student_map)
                    
                    # Save attachment file
                    filename = secure_filename(attachment_file.filename)
                    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                    filename = f"{timestamp}_{filename}"
                    attachment_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'assignments')
                    os.makedirs(attachment_dir, exist_ok=True)
                    file_path = os.path.join(attachment_dir, filename)
                    attachment_file.save(file_path)
                    # Use forward slash for web URLs (works on all platforms)
                    attachment_path = 'assignments/' + filename
                    attachment_filename = attachment_file.filename

            # Use ActivityService to create activity
            from services.activity_service import ActivityService
            new_activity = ActivityService.create_new_activity(
                instructor_id=current_user.id,
                title=title,
                activity_type=activity_type,
                description=description,
                quiz_category=quiz_category,
                due_date=due_date,
                student_id=student_id,
                course_ids=course_ids_list,
                attachment_path=attachment_path,
                attachment_filename=attachment_filename
            )
            
            student_name = "all students" if student_id is None else User.query.get(student_id).username
            course_names = ", ".join([c.name for c in valid_courses])
            flash(f'Assignment created successfully for {student_name} in {course_names}.', 'success')
            return redirect(url_for('instructor_assignments'))

        return render_template('instructor_assignment_create.html', 
                             students=all_students, 
                             courses=instructor_courses,
                             course_student_map=course_student_map)

    @app.route('/speaking', methods=['GET', 'POST'])
    @role_required('Student')
    def speaking():
        activity_id = request.args.get('activity_id')
        
        if request.method == 'POST':
            audio_file = request.files.get('audio_file')
            
            if not audio_file or audio_file.filename == '':
                flash("Please record or upload an audio file.", "danger")
                return redirect(url_for('speaking'))
            
            # Validate file format
            if not SubmissionService.validate_file_format(audio_file.filename, 'SPEAKING'):
                flash("Invalid audio format. Please upload MP3 or WAV files.", "danger")
                return redirect(url_for('speaking'))
            
            # Check file size (max 10MB)
            audio_file.seek(0, os.SEEK_END)
            file_size = audio_file.tell()
            audio_file.seek(0)
            
            if file_size > 10 * 1024 * 1024:  # 10MB
                flash("File size exceeds 10MB limit.", "danger")
                return redirect(url_for('speaking'))
            
            # Save audio file
            filename = secure_filename(audio_file.filename)
            # Add timestamp to avoid conflicts
            from datetime import datetime
            timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S_')
            filename = timestamp + filename
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            audio_file.save(file_path)
            
            # Check if already submitted for this activity
            if activity_id:
                existing = Submission.query.filter_by(
                    student_id=current_user.id,
                    activity_id=activity_id
                ).first()
                if existing:
                    flash("You have already submitted this assignment.", "danger")
                    return redirect(url_for('speaking'))
                
                # Check due date
                activity = LearningActivity.query.get(activity_id)
                if activity and activity.due_date:
                    if datetime.utcnow() > activity.due_date:
                        flash("This assignment has expired. The due date has passed.", "danger")
                        return redirect(url_for('speaking'))
            
            # Create submission
            new_sub, error_msg = SubmissionService.save_submission_text(
                student_id=current_user.id,
                activity_id=activity_id,
                submission_type='SPEAKING',
                text_content=None,
                file_path=filename
            )
            
            if not new_sub:
                flash(error_msg or "Failed to create submission.", "danger")
                return redirect(url_for('speaking'))
            
            # Check if AI is enabled
            if not AIService._is_ai_enabled():
                flash("AI features are currently disabled by administrator.", "danger")
                return redirect(url_for('speaking'))
            
            # Analyze with AI
            print(f"Starting AI analysis for speaking submission {new_sub.id}")
            ai_res = AIService.evaluate_speaking(file_path)
            
            # Process evaluation using GradingService
            if ai_res and ai_res.get('pronunciation_score') is not None:
                success = GradingService.process_speaking_evaluation(new_sub.id, ai_res)
                if success:
                    # Update goal progress for Speaking category
                    GoalService.update_goal_progress(current_user.id, 'Speaking')
                    NotificationService.notify_grade_ready(current_user.id, new_sub.id)
                    flash("Speaking analyzed successfully!", "success")
                    # Redirect to show results with submission_id parameter
                    return redirect(url_for('speaking', submission_id=new_sub.id))
                else:
                    flash("Failed to save grade.", "danger")
                    return redirect(url_for('speaking'))
            else:
                error_msg = ai_res.get('feedback', 'Unknown error') if ai_res else 'No response from AI'
                flash(f"Analysis failed: {error_msg}", "danger")
                return redirect(url_for('speaking'))
        
        # GET request - display page
        submission_id = request.args.get('submission_id', type=int)
        analysis_results = None
        
        # If viewing a specific submission, get its results
        if submission_id:
            submission = Submission.query.filter_by(id=submission_id, student_id=current_user.id).first()
            if submission and submission.grade:
                # Get tips from general_feedback or generate based on scores
                tips = []
                if submission.grade.pronunciation_score and submission.grade.pronunciation_score < 80:
                    tips.append("Practice difficult words slowly, then gradually increase speed")
                    tips.append("Record yourself and compare with native speakers")
                if submission.grade.fluency_score and submission.grade.fluency_score < 80:
                    tips.append("Read aloud daily to improve speech flow")
                    tips.append("Practice speaking without long pauses")
                if not tips:
                    tips.append("Keep up the excellent work!")
                    tips.append("Continue practicing to maintain your level")
                
                analysis_results = {
                    'pronunciation_score': submission.grade.pronunciation_score,
                    'fluency_score': submission.grade.fluency_score,
                    'feedback': submission.grade.general_feedback,
                    'tips': tips
                }
        
        # Get speaking submissions for stats
        submissions = Submission.query.filter_by(student_id=current_user.id, submission_type='SPEAKING').all()
        speaking_subs = [s for s in submissions if s.grade]
        
        # Calculate average score
        avg_score = 0.0
        if speaking_subs:
            scores = []
            for sub in speaking_subs:
                if sub.grade.pronunciation_score and sub.grade.fluency_score:
                    scores.append((sub.grade.pronunciation_score + sub.grade.fluency_score) / 2)
            avg_score = round(sum(scores) / len(scores), 1) if scores else 0.0
        
        # Get last practice date
        last_practice = None
        if speaking_subs:
            last_sub = max(speaking_subs, key=lambda x: x.created_at)
            last_practice = last_sub.created_at.strftime('%b %d') if last_sub.created_at else None
        
        total_recordings = len(speaking_subs)
        
        return render_template('speaking.html', 
                               avg_score=avg_score,
                               last_practice=last_practice,
                               total_recordings=total_recordings,
                               analysis_results=analysis_results)

    @app.route('/quizzes')
    @login_required
    def quizzes():
        # Get quizzes using QuizRepository
        user_quizzes = QuizRepository.get_quizzes(user_id=current_user.id)
        # Get available questions for new quizzes
        available_questions = Question.query.all()
        return render_template('quizzes.html', quizzes=user_quizzes, available_questions=available_questions)

    @app.route('/quiz/start', methods=['GET', 'POST'])
    @login_required
    def start_quiz():
        activity_id = request.args.get('activity_id', type=int)
        category = None
        
        # Get category from POST request if available
        if request.method == 'POST':
            category = request.form.get('category')
            if category and category.strip() == '':
                category = None
            # Keep 'mixed' as 'mixed' - don't convert to None (needed for filtering)
            # Note: For question fetching, 'mixed' means all categories, but we still save it as 'mixed'
        
        # If started from assignment, get category from activity and verify student access
        if activity_id:
            activity = LearningActivity.query.get(activity_id)
            if activity:
                # Verify student has access to this activity
                if current_user.role == 'Student':
                    if activity.student_id is not None and activity.student_id != current_user.id:
                        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                            from flask import jsonify
                            return jsonify({'success': False, 'message': 'You do not have access to this assignment.'}), 403
                        flash("You do not have access to this assignment.", "danger")
                        return redirect(url_for('assignments'))
                
                if activity.activity_type == 'QUIZ':
                    category = activity.quiz_category
        
        # Check if questions are available before attempting to get them
        available, message = QuizService.check_questions_available(category)
        if not available:
            # Check if this is an AJAX request
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                from flask import jsonify
                return jsonify({'success': False, 'message': message}), 400
            flash(message, "danger")
            return redirect(url_for('quizzes'))
        
        # Get questions using QuizService with optional category
        # For 'mixed', pass None to get questions from all categories
        question_category = None if category == 'mixed' else category
        questions = QuizService.get_questions(limit=5, category=question_category)
        
        if not questions:
            # This shouldn't happen if check_questions_available passed, but handle it anyway
            error_msg = f"No questions found for category '{category}'." if category else "No questions found."
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                from flask import jsonify
                return jsonify({'success': False, 'message': error_msg}), 400
            flash(error_msg, "danger")
            return redirect(url_for('quizzes'))
        
        # Store questions in session for quiz flow
        from flask import session
        session['quiz_questions'] = [q.id for q in questions]
        session['quiz_answers'] = {}
        session['quiz_current'] = 0
        session['quiz_started'] = True
        session['quiz_category'] = category  # Store category for later use
        if activity_id:
            session['quiz_activity_id'] = activity_id
        
        # Check if this is an AJAX request
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            from flask import jsonify
            return jsonify({
                'success': True,
                'redirect': url_for('quiz_question')
            })
        
        return redirect(url_for('quiz_question'))

    @app.route('/quiz/question', methods=['GET'])
    @login_required
    def quiz_question():
        """Display the dynamic quiz interface"""
        from flask import session
        
        if not session.get('quiz_started'):
            flash("Please start a quiz first.", "danger")
            return redirect(url_for('quizzes'))
        
        # The new dynamic interface loads questions via JavaScript API
        # So we just render the template
        return render_template('quiz_question.html')

    @app.route('/quiz/questions', methods=['GET'])
    @login_required
    def get_quiz_questions():
        """API endpoint to fetch all quiz questions as JSON"""
        from flask import session, jsonify
        
        if not session.get('quiz_started'):
            return jsonify({'error': 'No quiz started'}), 400
        
        question_ids = session.get('quiz_questions', [])
        questions = []
        
        for q_id in question_ids:
            question = Question.query.get(q_id)
            if question:
                questions.append({
                    'id': question.id,
                    'question_text': question.question_text,
                    'option_a': question.option_a,
                    'option_b': question.option_b,
                    'option_c': question.option_c,
                    'option_d': question.option_d,
                    'correct_answer': question.correct_answer,
                    'category': question.category
                })
        
        return jsonify({
            'questions': questions,
            'total': len(questions)
        })
    
    @app.route('/quiz/submit', methods=['POST'])
    @login_required
    def submit_quiz():
        """API endpoint to submit quiz answers and calculate score"""
        from flask import session, jsonify
        
        if not session.get('quiz_started'):
            return jsonify({'error': 'No quiz started'}), 400
        
        data = request.get_json()
        answers = data.get('answers', {})  # {question_id: answer}
        time_spent = data.get('time_spent', 0)  # in seconds
        
        question_ids = session.get('quiz_questions', [])
        
        # Convert answers to string keys format (for compatibility)
        answers_dict = {str(q_id): answers.get(str(q_id), '') for q_id in question_ids}
        
        # Calculate score using QuizService
        correct, total, score = QuizService.calculate_final_score(question_ids, answers_dict)

        # Build per-question details for result page with AI explanations
        from services.ai_service import AIService
        import concurrent.futures
        import threading
        
        details = []
        incorrect_questions = []  # Store questions that need AI explanations
        
        for q_id in question_ids:
            question = Question.query.get(q_id)
            user_answer = answers.get(str(q_id), '')
            correct_answer = question.correct_answer if question else None
            is_correct = user_answer and question and user_answer.upper() == correct_answer.upper()
            
            detail_item = {
                'question_text': question.question_text if question else '',
                'user_answer': user_answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
                'explanation': None  # Will be populated by AI
            }
            
            # If incorrect, add to list for AI explanation generation
            if not is_correct and question and user_answer:
                incorrect_questions.append({
                    'detail_item': detail_item,
                    'question_text': question.question_text,
                    'user_answer': user_answer,
                    'correct_answer': correct_answer
                })
            
            details.append(detail_item)
        
        # Generate AI explanations for incorrect answers (with timeout for performance)
        if incorrect_questions:
            # Check if AI is enabled
            ai_enabled = AIService._is_ai_enabled()
            
            if not ai_enabled:
                # AI is disabled - set placeholder message for all incorrect questions
                for item in incorrect_questions:
                    item['detail_item']['explanation'] = "AI is disabled by admin."
            else:
                def generate_explanation(item):
                    try:
                        explanation = AIService.generate_quiz_explanation(
                            item['question_text'],
                            item['user_answer'],
                            item['correct_answer']
                        )
                        item['detail_item']['explanation'] = explanation
                    except Exception as e:
                        print(f"Error generating explanation: {e}")
                        item['detail_item']['explanation'] = "Generating AI analysis... Please try again later."
                
                # Use ThreadPoolExecutor for parallel processing (meets NFR1: 10-second response time)
                with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                    futures = {executor.submit(generate_explanation, item): item for item in incorrect_questions}
                    # Wait for all with timeout (max 8 seconds to leave buffer for other operations)
                    done, not_done = concurrent.futures.wait(futures.keys(), timeout=8.0)
                    
                    # For any that didn't complete, set a placeholder
                    for future in not_done:
                        future.cancel()
                        # Find the corresponding item and set placeholder
                        item = futures[future]
                        if item['detail_item']['explanation'] is None:
                            item['detail_item']['explanation'] = "Generating AI analysis... Please refresh the page in a moment."
        
        # Save quiz result and detailed answers using QuizService
        quiz_category = session.get('quiz_category')  # Get category from session
        
        # If category is None, try to determine from questions
        if quiz_category is None:
            # Get first question to determine category
            if question_ids:
                first_question = Question.query.get(question_ids[0])
                if first_question and first_question.category:
                    quiz_category = first_question.category.lower()
        
        quiz_title = "Grammar Quiz"
        
        # Determine quiz title and category based on assignment or standalone quiz
        if quiz_category:
            # Map category to title
            category_titles = {
                'grammar': 'Grammar Quiz',
                'vocabulary': 'Vocabulary Quiz',
                'reading': 'Reading Quiz',
                'mixed': 'Mixed Quiz'
            }
            quiz_title = category_titles.get(quiz_category.lower(), 'Quiz')
        else:
            # Default to Grammar if no category found
            quiz_category = 'grammar'
            quiz_title = 'Grammar Quiz'
        
        # Check if this was an assignment submission
        activity_id = session.get('quiz_activity_id')
        if activity_id:
            activity = LearningActivity.query.get(activity_id)
            if activity:
                quiz_title = activity.title
                # Use activity's category if available
                if activity.quiz_category:
                    quiz_category = activity.quiz_category
                
                # Create Submission
                new_sub = Submission(
                    student_id=current_user.id,
                    activity_id=activity.id,
                    submission_type='QUIZ',
                    text_content=f"Quiz completed: {activity.title} (Score: {score}%)"
                )
                db.session.add(new_sub)
                db.session.flush() # get id
                
                # Create Grade
                # Quiz grades are auto-approved since they're automatically graded
                new_grade = Grade(
                    submission_id=new_sub.id,
                    score=score,
                    general_feedback=f"Auto-graded quiz. Correct: {correct}/{total}",
                    instructor_approved=True  # Auto-approved for quizzes
                )
                db.session.add(new_grade)
                db.session.commit() # Commit submission and grade
        
        QuizService.save_result(current_user.id, quiz_title, score, details=details, category=quiz_category)
        
        # Clear session
        session.pop('quiz_started', None)
        session.pop('quiz_questions', None)
        session.pop('quiz_answers', None)
        session.pop('quiz_current', None)
        session.pop('quiz_activity_id', None)
        session.pop('quiz_category', None)  # Clear category from session
        
        return jsonify({
            'success': True,
            'score': score,
            'correct': correct,
            'total': total,
            'time_spent': time_spent,
            'redirect': url_for('quiz_result', score=score, correct=correct, total=total)
        })
    
    @app.route('/quiz/result')
    @login_required
    def quiz_result():
        """Display quiz result page"""
        score = request.args.get('score', type=float, default=0)
        correct = request.args.get('correct', type=int, default=0)
        total = request.args.get('total', type=int, default=0)
        
        # Fetch the most recent quiz for this user to get details
        details = []
        if total > 0:
            recent_quiz = Quiz.query.filter_by(user_id=current_user.id).order_by(Quiz.date_taken.desc()).first()
            if recent_quiz:
                detail_rows = QuizDetail.query.filter_by(quiz_id=recent_quiz.id).all()
                for d in detail_rows:
                    details.append({
                        'question_text': d.question_text,
                        'user_answer': d.user_answer,
                        'correct_answer': d.correct_answer,
                        'is_correct': d.is_correct,
                        'explanation': d.explanation if hasattr(d, 'explanation') else None,  # Get AI-generated explanation
                    })
        
        return render_template('quiz_result.html', 
                             score=score, 
                             correct=correct, 
                             total=total, 
                             details=details, 
                             is_review=False)

    @app.route('/quiz/finish', methods=['GET', 'POST'])
    @login_required
    def finish_quiz():
        from flask import session
        
        if not session.get('quiz_started'):
            flash("Please start a quiz first.", "danger")
            return redirect(url_for('quizzes'))
        
        question_ids = session.get('quiz_questions', [])
        answers = session.get('quiz_answers', {})
        
        # Calculate score using QuizService
        correct, total, score = QuizService.calculate_final_score(question_ids, answers)

        # Build per-question details for result page
        details = []
        for q_id in question_ids:
            question = Question.query.get(q_id)
            user_answer = answers.get(str(q_id))
            correct_answer = question.correct_answer if question else None
            is_correct = user_answer and question and user_answer.upper() == correct_answer.upper()
            details.append({
                'question_text': question.question_text if question else '',
                'user_answer': user_answer,
                'correct_answer': correct_answer,
                'is_correct': is_correct,
            })
        
        # Save quiz result and detailed answers using QuizService
        quiz_category = session.get('quiz_category')  # Get category from session
        
        # If category is None, try to determine from questions
        if quiz_category is None:
            # Get first question to determine category
            if question_ids:
                first_question = Question.query.get(question_ids[0])
                if first_question and first_question.category:
                    quiz_category = first_question.category.lower()
        
        quiz_title = "Grammar Quiz"
        
        # Determine quiz title and category based on assignment or standalone quiz
        if quiz_category:
            # Map category to title
            category_titles = {
                'grammar': 'Grammar Quiz',
                'vocabulary': 'Vocabulary Quiz',
                'reading': 'Reading Quiz',
                'mixed': 'Mixed Quiz'
            }
            quiz_title = category_titles.get(quiz_category.lower(), 'Quiz')
        else:
            # Default to Grammar if no category found
            quiz_category = 'grammar'
            quiz_title = 'Grammar Quiz'
        
        # Check if this was an assignment submission
        activity_id = session.get('quiz_activity_id')
        if activity_id:
            activity = LearningActivity.query.get(activity_id)
            if activity:
                quiz_title = activity.title
                # Use activity's category if available
                if activity.quiz_category:
                    quiz_category = activity.quiz_category
                
                # Create Submission
                new_sub = Submission(
                    student_id=current_user.id,
                    activity_id=activity.id,
                    submission_type='QUIZ',
                    text_content=f"Quiz completed: {activity.title} (Score: {score}%)"
                )
                db.session.add(new_sub)
                db.session.flush() # get id
                
                # Create Grade
                # Quiz grades are auto-approved since they're automatically graded
                new_grade = Grade(
                    submission_id=new_sub.id,
                    score=score,
                    general_feedback=f"Auto-graded quiz. Correct: {correct}/{total}",
                    instructor_approved=True  # Auto-approved for quizzes
                )
                db.session.add(new_grade)
                db.session.commit() # Commit submission and grade
                flash("Assignment marked as completed!", "success")
        
        QuizService.save_result(current_user.id, quiz_title, score, details=details, category=quiz_category)
        
        # Update goal progress for Quiz category
        GoalService.update_goal_progress(current_user.id, 'Quiz')
        
        # Clear session
        session.pop('quiz_started', None)
        session.pop('quiz_questions', None)
        session.pop('quiz_answers', None)
        session.pop('quiz_current', None)
        session.pop('quiz_activity_id', None)
        session.pop('quiz_category', None)  # Clear category from session
        
        return render_template('quiz_result.html', score=score, correct=correct, total=total, details=details, is_review=False)

    @app.route('/quiz/review/<int:quiz_id>')
    @login_required
    def quiz_review(quiz_id):
        quiz = Quiz.query.get_or_404(quiz_id)

        # Permission: student can only see own quiz, instructor can see all
        if current_user.role != 'Instructor' and quiz.user_id != current_user.id:
            flash("You don't have permission to view this quiz.", "danger")
            return redirect(url_for('dashboard'))

        details_rows = QuizDetail.query.filter_by(quiz_id=quiz.id).all()

        # Fallback: if no stored details, just show simple result
        details = []
        correct = 0
        total = 0
        if details_rows:
            for d in details_rows:
                details.append({
                    'question_text': d.question_text,
                    'user_answer': d.user_answer,
                    'correct_answer': d.correct_answer,
                    'is_correct': d.is_correct,
                })
                total += 1
                if d.is_correct:
                    correct += 1

        return render_template(
            'quiz_result.html',
            score=quiz.score,
            correct=correct,
            total=total,
            details=details,
            is_review=True
        )

    @app.route('/goals', methods=['GET', 'POST'])
    @login_required
    @role_required('Student')
    def goals():
        if request.method == 'POST':
            try:
                # Verify user is logged in
                if not current_user or not current_user.is_authenticated:
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'error', 'success': False, 'message': 'User not authenticated. Please log in again.'}), 401
                    flash('Please log in to create goals.', 'error')
                    return redirect(url_for('login'))
                
            # Handle goal creation
                title = request.form.get('goal_name') or request.form.get('title')
                category = request.form.get('category')
                target_score_raw = request.form.get('target_value') or request.form.get('target_score')
                target_date_str = request.form.get('target_date')
                
                # Convert to float - handle comma as decimal separator (European format)
                target_score = None
                if target_score_raw:
                    try:
                        # Replace comma with dot for European decimal format
                        target_score_str = str(target_score_raw).replace(',', '.')
                        target_score = float(target_score_str)
                    except (ValueError, TypeError) as e:
                        print(f"Error converting target_score to float: {str(e)}, value: {target_score_raw}")
                        target_score = None
                
                # Parse target date - prioritize European format (DD.MM.YYYY)
                target_date_obj = None
                if target_date_str:
                    date_str = str(target_date_str).strip()
                    if date_str:
                        try:
                            # Try DD.MM.YYYY format first (European format)
                            target_date_obj = datetime.strptime(date_str, '%d.%m.%Y')
                        except (ValueError, TypeError):
                            try:
                                # Try standard HTML date format (YYYY-MM-DD)
                                target_date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                            except (ValueError, TypeError):
                                try:
                                    # Try DD/MM/YYYY format
                                    target_date_obj = datetime.strptime(date_str, '%d/%m/%Y')
                                except (ValueError, TypeError):
                                    print(f"Warning: Could not parse date format: {date_str}")
                                    target_date_obj = None
                
                # Validate target date is not in the past
                if target_date_obj:
                    today = datetime.now().date()
                    target_date_only = target_date_obj.date() if hasattr(target_date_obj, 'date') else target_date_obj
                    if target_date_only < today:
                        error_msg = 'Target date cannot be in the past. Please select today or a future date.'
                        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                            return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 400
                        flash(error_msg, 'error')
                        return redirect(url_for('goals'))
                
                # Validate required fields
                if not title or not title.strip():
                    error_msg = 'Goal name is required.'
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 400
                    flash(error_msg, 'error')
                    return redirect(url_for('goals'))
                
                if not category:
                    error_msg = 'Category is required.'
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 400
                    flash(error_msg, 'error')
                    return redirect(url_for('goals'))
                
                if target_score is None:
                    error_msg = 'Target score is required.'
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 400
                    flash(error_msg, 'error')
                    return redirect(url_for('goals'))
                
                # Validate category
                valid_categories = ['Writing', 'Speaking', 'Quiz', 'Grammar', 'Vocabulary', 'Reading', 'Overall']
                if category not in valid_categories:
                    error_msg = f'Invalid category. Must be one of: {", ".join(valid_categories)}.'
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 400
                    flash(error_msg, 'error')
                    return redirect(url_for('goals'))
                
                # Validate target score range
                if target_score < 0 or target_score > 100:
                    error_msg = 'Target score must be between 0 and 100.'
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 400
                    flash(error_msg, 'error')
                    return redirect(url_for('goals'))
                
                # Check for duplicate active goals in same category
                existing_goals = GoalService.get_user_goals(current_user.id)
                duplicate = any(g.category == category and g.status == 'In Progress' for g in existing_goals)
                if duplicate:
                    error_msg = f'You already have an active goal for {category}. Please complete or delete it first.'
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 400
                    flash(error_msg, 'warning')
                    return redirect(url_for('goals'))
                
                # Use GoalService to create goal (current_score is auto-calculated, starts at 0)
                new_goal = GoalService.create_goal(
                    user_id=current_user.id,
                    title=title.strip(),
                    category=category,
                    target_score=float(target_score),
                    current_score=0.0,  # Always start at 0, auto-updated
                    target_date=target_date_obj
                )
                
                # Return JSON for AJAX requests
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'status': 'success', 'success': True, 'message': 'Goal added successfully!'}), 200
                
                flash('Goal added successfully!', 'success')
                return redirect(url_for('goals'))
                
            except Exception as e:
                # Advanced error logging with traceback
                error_traceback = traceback.format_exc()
                error_message = str(e)
                error_type = type(e).__name__
                
                print("=" * 80)
                print("ERROR CREATING GOAL - Full Traceback:")
                print(error_traceback)
                print("=" * 80)
                print(f"Error message: {error_message}")
                print(f"Error type: {error_type}")
                print(f"User ID: {current_user.id if current_user and current_user.is_authenticated else 'Not authenticated'}")
                print(f"Request data: title={request.form.get('goal_name')}, category={request.form.get('category')}, target_score={request.form.get('target_value')}, target_date={request.form.get('target_date')}")
                print("=" * 80)
                
                # Always return JSON for AJAX requests, even on error
                # Return exact error message so user can see it in browser
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({
                        'status': 'error',
                        'success': False,
                        'message': f'Error creating goal: {error_message}',
                        'error_type': error_type,
                        'details': 'Check server logs for full traceback.'
                    }), 500
                flash(f'Error creating goal: {error_message}', 'error')
                return redirect(url_for('goals'))
        
        # GET request - display goals using GoalService
        # Note: datetime is already imported at the top of the file
        user_goals = GoalService.get_user_goals(current_user.id)
        goals_summary = GoalService.get_goals_summary(current_user.id)
        return render_template('goals.html', goals=user_goals, goals_summary=goals_summary, now=datetime.utcnow())
    
    @app.route('/goals/<int:goal_id>', methods=['GET', 'PUT'])
    @login_required
    @role_required('Student')
    def get_or_update_goal(goal_id):
        try:
            # Verify user is logged in
            if not current_user or not current_user.is_authenticated:
                return jsonify({'status': 'error', 'success': False, 'message': 'User not authenticated. Please log in again.'}), 401
            
            goal = GoalRepository.get_goal_by_id(goal_id)
            
            if not goal:
                return jsonify({'status': 'error', 'success': False, 'message': 'Goal not found'}), 404
            
            # Ensure user can only access their own goals
            if goal.user_id != current_user.id:
                return jsonify({'status': 'error', 'success': False, 'message': 'Permission denied'}), 403
            
            if request.method == 'GET':
                # Return goal data
                return jsonify({
                    'status': 'success',
                    'success': True,
                    'goal': {
                        'id': goal.id,
                        'goal_name': goal.title,  # For backward compatibility
                        'title': goal.title,
                        'category': goal.category,
                        'target_value': goal.target_score,  # For backward compatibility
                        'target_score': goal.target_score,
                        'current_value': goal.current_score,  # For backward compatibility
                        'current_score': goal.current_score,
                        'status': goal.status,
                        'target_date': goal.target_date.isoformat() if goal.target_date else None
                    }
                })
            
            elif request.method == 'PUT':
                # Update goal
                title = request.form.get('goal_name') or request.form.get('title')
                target_score_raw = request.form.get('target_value') or request.form.get('target_score')
                category = request.form.get('category')
                status = request.form.get('status')
                target_date_str = request.form.get('target_date')
                
                # Convert to float - handle comma as decimal separator (European format)
                target_score = None
                if target_score_raw:
                    try:
                        # Replace comma with dot for European decimal format
                        target_score_str = str(target_score_raw).replace(',', '.')
                        target_score = float(target_score_str)
                    except (ValueError, TypeError):
                        target_score = None
                
                # Parse target date - prioritize European format (DD.MM.YYYY)
                target_date_obj = None
                if target_date_str:
                    date_str = str(target_date_str).strip()
                    if date_str:
                        try:
                            # Try DD.MM.YYYY format first (European format)
                            target_date_obj = datetime.strptime(date_str, '%d.%m.%Y')
                        except (ValueError, TypeError):
                            try:
                                # Try standard HTML date format (YYYY-MM-DD)
                                target_date_obj = datetime.strptime(date_str, '%Y-%m-%d')
                            except (ValueError, TypeError):
                                try:
                                    # Try DD/MM/YYYY format
                                    target_date_obj = datetime.strptime(date_str, '%d/%m/%Y')
                                except (ValueError, TypeError):
                                    print(f"Warning: Could not parse date format: {date_str}")
                                    target_date_obj = None
                
                # Validate target date is not in the past
                if target_date_obj:
                    today = datetime.now().date()
                    target_date_only = target_date_obj.date() if hasattr(target_date_obj, 'date') else target_date_obj
                    if target_date_only < today:
                        error_msg = 'Target date cannot be in the past. Please select today or a future date.'
                        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                            return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 400
                        flash(error_msg, 'error')
                        return redirect(url_for('goals'))
                
                # Don't allow updating current_score manually - it's auto-updated
                updated_goal = GoalService.update_goal(
                    goal_id=goal_id,
                    title=title,
                    category=category,
                    target_score=target_score,
                    status=status,
                    target_date=target_date_obj
                )
                
                if updated_goal:
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'success', 'success': True, 'message': 'Goal updated successfully!'}), 200
                    flash('Goal updated successfully!', 'success')
                    return redirect(url_for('goals'))
                else:
                    error_msg = 'Failed to update goal. Goal not found.'
                    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                        return jsonify({'status': 'error', 'success': False, 'message': error_msg}), 404
                    flash(error_msg, 'error')
                    return redirect(url_for('goals'))

        except Exception as e:
            # Advanced error logging with traceback
            error_traceback = traceback.format_exc()
            print("=" * 80)
            print("ERROR IN get_or_update_goal - Full Traceback:")
            print(error_traceback)
            print("=" * 80)
            print(f"Error message: {str(e)}")
            print(f"Error type: {type(e).__name__}")
            print(f"Goal ID: {goal_id}")
            print(f"User ID: {current_user.id if current_user and current_user.is_authenticated else 'Not authenticated'}")
            print("=" * 80)
            
            # Always return JSON for AJAX requests, even on error
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({
                    'status': 'error',
                    'success': False,
                    'message': f'Error processing goal request: {str(e)}. Please check the server logs for details.'
                }), 500
            flash(f'Error: {str(e)}', 'error')
            return redirect(url_for('goals'))

    @app.route('/goals/<int:goal_id>/complete', methods=['POST'])
    @login_required
    @role_required('Student')
    def mark_goal_completed(goal_id):
        goal = GoalRepository.get_goal_by_id(goal_id)
        
        if not goal:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'message': 'Goal not found'}), 404
            flash('Goal not found.', 'error')
            return redirect(url_for('goals'))
        
        # Ensure user can only complete their own goals
        if goal.user_id != current_user.id:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'message': 'Permission denied'}), 403
            flash('You do not have permission to complete this goal.', 'error')
            return redirect(url_for('goals'))
        
        try:
            # Use GoalService to mark as completed
            updated_goal = GoalService.mark_as_completed(goal_id)
            
            if updated_goal:
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': True, 'message': 'Goal marked as completed!'}), 200
                flash('Goal marked as completed!', 'success')
                return redirect(url_for('goals'))
            else:
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': False, 'message': 'Failed to complete goal'}), 500
                flash('Failed to complete goal.', 'error')
                return redirect(url_for('goals'))
        except Exception as e:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'message': str(e)}), 500
            flash(f'Error: {str(e)}', 'error')
            return redirect(url_for('goals'))

    @app.route('/delete-goal/<int:goal_id>', methods=['POST'])
    @login_required
    @role_required('Student')
    def delete_goal(goal_id):
        try:
            goal = GoalRepository.get_goal_by_id(goal_id)
            
            if not goal:
                error_msg = 'Goal not found'
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': False, 'message': error_msg}), 404
                flash(error_msg, 'error')
                return redirect(url_for('goals'))
            
            # Ensure user can only delete their own goals
            if goal.user_id != current_user.id:
                error_msg = 'Permission denied'
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': False, 'message': error_msg}), 403
                flash('You do not have permission to delete this goal.', 'error')
                return redirect(url_for('goals'))
            
            # Use GoalService to delete goal
            success = GoalService.delete_goal(goal_id)
            
            if success:
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': True, 'message': 'Goal deleted successfully!'}), 200
                flash('Goal deleted successfully!', 'success')
                return redirect(url_for('goals'))
            else:
                error_msg = 'Failed to delete goal'
                if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                    return jsonify({'success': False, 'message': error_msg}), 500
                flash(error_msg, 'error')
                return redirect(url_for('goals'))
        except Exception as e:
            # Log error for debugging
            print(f"Error deleting goal: {str(e)}")
            # Always return JSON for AJAX requests, even on error
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'message': f'Error deleting goal: {str(e)}'}), 500
            flash(f'Error: {str(e)}', 'error')
            return redirect(url_for('goals'))
    
    @app.route('/profile')
    @login_required
    def profile():
        from models.entities import Course, Enrollment, LearningActivity, Submission, Grade, Quiz
        from sqlalchemy import func
        
        user = current_user
        stats = {}
        
        # Student-specific stats
        if user.role == 'Student':
            submissions = Submission.query.filter_by(student_id=user.id).all()
            quizzes = Quiz.query.filter_by(user_id=user.id).all()
            
            # Calculate streak
            current_streak = 0
            if submissions:
                from datetime import timedelta
                submission_dates = set(s.created_at.date() for s in submissions)
                today = datetime.utcnow().date()
                date = today
                while date in submission_dates:
                    current_streak += 1
                    date -= timedelta(days=1)
            
            # Calculate average score
            graded_subs = [s for s in submissions if s.grade]
            avg_score = 0.0
            if graded_subs:
                scores = []
                for sub in graded_subs:
                    if sub.submission_type == 'SPEAKING' and sub.grade.pronunciation_score and sub.grade.fluency_score:
                        scores.append((sub.grade.pronunciation_score + sub.grade.fluency_score) / 2)
                    elif sub.grade.score:
                        scores.append(sub.grade.score)
                avg_score = round(sum(scores) / len(scores), 1) if scores else 0.0
            
            stats = {
                'total_tasks': len(submissions),
                'avg_score': avg_score,
                'streak': current_streak,
                'completed_quizzes': len(quizzes),
                'total_submissions': len(submissions)
            }
        
        # Instructor-specific stats
        elif user.role == 'Instructor':
            courses = Course.query.filter_by(instructor_id=user.id).all()
            activities = LearningActivity.query.filter_by(instructor_id=user.id).all()
            
            # Count students taught (unique students across all courses)
            all_students = set()
            for course in courses:
                enrollments = Enrollment.query.filter_by(course_id=course.id, status='active').all()
                all_students.update(e.student_id for e in enrollments)
            
            # Count total submissions to instructor's activities
            total_submissions = Submission.query.filter(
                Submission.activity_id.in_([a.id for a in activities])
            ).count()
            
            # Count pending reviews
            pending_reviews = Submission.query.filter(
                Submission.activity_id.in_([a.id for a in activities]),
                ~Submission.id.in_(db.session.query(Grade.submission_id))
            ).count()
            
            # Average student score
            graded_submissions = Submission.query.filter(
                Submission.activity_id.in_([a.id for a in activities])
            ).join(Grade).all()
            
            avg_student_score = 0.0
            if graded_submissions:
                scores = []
                for sub in graded_submissions:
                    if sub.submission_type == 'SPEAKING' and sub.grade.pronunciation_score and sub.grade.fluency_score:
                        scores.append((sub.grade.pronunciation_score + sub.grade.fluency_score) / 2)
                    elif sub.grade.score:
                        scores.append(sub.grade.score)
                avg_student_score = round(sum(scores) / len(scores), 1) if scores else 0.0
            
            stats = {
                'courses_taught': len(courses),
                'students_taught': len(all_students),
                'assignments_created': len(activities),
                'total_submissions': total_submissions,
                'pending_reviews': pending_reviews,
                'avg_student_score': avg_student_score
            }
        
        # Admin-specific stats
        elif user.role == 'Admin':
            from services.admin_service import AdminService
            platform_stats = AdminService.get_user_statistics()
            
            stats = {
                'total_users': platform_stats['total_users'],
                'total_students': platform_stats['total_students'],
                'total_instructors': platform_stats['total_instructors'],
                'total_courses': platform_stats['total_courses'],
                'active_enrollments': platform_stats['active_enrollments']
            }
        
        return render_template('profile.html', stats=stats)
    
    @app.route('/update_bio', methods=['POST'])
    @login_required
    def update_bio():
        new_bio = request.form.get('new_bio', '').strip()
        try:
            # Check if User model has bio attribute, if not we'll add it dynamically
            if hasattr(current_user, 'bio'):
                current_user.bio = new_bio
            else:
                # Try to set it anyway - SQLAlchemy will handle if column doesn't exist
                setattr(current_user, 'bio', new_bio)
            db.session.commit()
            flash('Bio updated successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating bio: {str(e)}', 'danger')
        return redirect(url_for('profile'))
    
    @app.route('/update_personal_info', methods=['POST'])
    @login_required
    def update_personal_info():
        university = request.form.get('university', '').strip()
        grade = request.form.get('grade', '').strip()
        teacher = request.form.get('teacher', '').strip()
        phone = request.form.get('phone', '').strip()
        education_status = request.form.get('education_status', '').strip()
        
        try:
            user = current_user
            # Update fields if they exist in the model
            if hasattr(user, 'university'):
                user.university = university if university else None
            if hasattr(user, 'grade'):
                user.grade = grade if grade else None
            if hasattr(user, 'teacher'):
                user.teacher = teacher if teacher else None
            if hasattr(user, 'phone'):
                user.phone = phone if phone else None
            if hasattr(user, 'education_status'):
                user.education_status = education_status if education_status else None
            
            db.session.commit()
            flash('Personal information updated successfully!', 'success')
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating personal information: {str(e)}', 'danger')
        return redirect(url_for('profile'))
    
    @app.route('/upload_profile_picture', methods=['POST'])
    @login_required
    def upload_profile_picture():
        from werkzeug.utils import secure_filename
        import uuid
        
        if 'profile_image' not in request.files:
            return jsonify({'success': False, 'message': 'No file provided'}), 400
        
        file = request.files['profile_image']
        if file.filename == '':
            return jsonify({'success': False, 'message': 'No file selected'}), 400
        
        # Check if file is an image
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif'}
        if '.' in file.filename and file.filename.rsplit('.', 1)[1].lower() not in allowed_extensions:
            return jsonify({'success': False, 'message': 'Invalid file type. Only images are allowed.'}), 400
        
        try:
            # Generate unique filename
            file_ext = file.filename.rsplit('.', 1)[1].lower()
            filename = f"{current_user.id}_{uuid.uuid4().hex[:8]}.{file_ext}"
            
            # Save to profile_pics folder
            profile_pics_folder = os.path.join(os.path.abspath(os.path.dirname(__file__)), 'static/profile_pics')
            os.makedirs(profile_pics_folder, exist_ok=True)
            filepath = os.path.join(profile_pics_folder, filename)
            file.save(filepath)
            
            # Update user profile_image
            old_filename = current_user.profile_image if hasattr(current_user, 'profile_image') else None
            
            # Try to set profile_image attribute
            if hasattr(current_user, 'profile_image'):
                current_user.profile_image = filename
            else:
                setattr(current_user, 'profile_image', filename)
            
            db.session.commit()
            
            # Delete old profile picture if exists
            if old_filename:
                old_filepath = os.path.join(profile_pics_folder, old_filename)
                if os.path.exists(old_filepath):
                    try:
                        os.remove(old_filepath)
                    except:
                        pass
            
            return jsonify({'success': True, 'message': 'Profile picture uploaded successfully!'})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'message': f'Error uploading picture: {str(e)}'}), 500
    
    @app.route('/settings')
    @login_required
    def settings():
        # Settings page removed - redirect to profile
        return redirect(url_for('profile'))
    @app.route('/export/pdf')
    @login_required
    def export_pdf():
        """Export student submissions to PDF"""
        # Check if reportlab is available
        if not REPORTLAB_AVAILABLE:
            return "PDF generation requires reportlab library. Please install with: pip install reportlab", 500
        
        try:
            from models.entities import Submission
            
            # Get all submissions for the current user, sorted chronologically (oldest first)
            submissions = db.session.query(Submission).filter_by(student_id=current_user.id).order_by(Submission.created_at.asc()).all()
            
            # Create PDF buffer
            buffer = io.BytesIO()
            try:
                p = canvas.Canvas(buffer, pagesize=letter)
                width, height = letter
                
                # Define consistent font sizes for entire document
                TITLE_FONT_SIZE = 18
                HEADER_INFO_FONT_SIZE = 11
                TABLE_HEADER_FONT_SIZE = 11
                TABLE_DATA_FONT_SIZE = 10
                
                # Header section
                p.setFont("Helvetica-Bold", TITLE_FONT_SIZE)
                p.drawString(100, height - 50, "AAFS AI - Academic Progress Report")
                p.setFont("Helvetica", HEADER_INFO_FONT_SIZE)
                p.drawString(100, height - 80, f"Student: {current_user.username}")
                # Use GMT+3 for generated time
                generated_time = get_gmt3_now()
                p.drawString(100, height - 100, f"Generated: {generated_time.strftime('%Y-%m-%d %H:%M')} (GMT+3)")
                p.line(100, height - 110, 500, height - 110)
                
                if submissions:
                    
                    # Table headers - adjusted column positions
                    y = height - 130
                    p.setFont("Helvetica-Bold", TABLE_HEADER_FONT_SIZE)
                    p.drawString(100, y, "Submission Date")
                    p.drawString(280, y, "Type")
                    p.drawString(380, y, "Score")
                    p.line(100, y - 5, 500, y - 5)
                    
                    # Table rows
                    p.setFont("Helvetica", TABLE_DATA_FONT_SIZE)
                    y -= 25
                    
                    for sub in submissions:
                        # Check if we need a new page
                        if y < 100:
                            p.showPage()
                            # Re-draw headers on new page with same font sizes
                            y = height - 130
                            p.setFont("Helvetica-Bold", TABLE_HEADER_FONT_SIZE)
                            p.drawString(100, y, "Submission Date")
                            p.drawString(280, y, "Type")
                            p.drawString(380, y, "Score")
                            p.line(100, y - 5, 500, y - 5)
                            # Set data font for rows
                            p.setFont("Helvetica", TABLE_DATA_FONT_SIZE)
                            y -= 25
                        
                        try:
                            # Get score safely
                            if sub.grade and sub.grade.score is not None:
                                score = sub.grade.score
                            else:
                                score = 'N/A'
                            
                            # Format submission date with time in GMT+3
                            if sub.created_at:
                                sub_date_gmt3 = utc_to_gmt3(sub.created_at)
                                date_str = sub_date_gmt3.strftime('%Y-%m-%d %H:%M')
                            else:
                                date_str = 'N/A'
                            
                            # Format type safely
                            if sub.submission_type:
                                type_str = sub.submission_type.capitalize()
                            else:
                                type_str = 'Unknown'
                            
                            # Draw row
                            p.drawString(100, y, date_str)
                            p.drawString(280, y, type_str)
                            p.drawString(380, y, f"{score}%" if isinstance(score, (int, float)) else str(score))
                            
                            y -= 20
                        except Exception as e:
                            # Skip problematic submissions and continue
                            continue
                else:
                    # No submissions message
                    y = height - 130
                    p.setFont("Helvetica", HEADER_INFO_FONT_SIZE)
                    p.drawString(100, y, "No submissions found.")
                
                # Finalize and save PDF
                p.showPage()
                p.save()
            finally:
                # Ensure buffer is ready for reading
                buffer.seek(0)
            
            # Generate filename with GMT+3 date (safe ASCII only)
            filename = f"academic_report_{current_user.id}_{get_gmt3_now().strftime('%Y%m%d')}.pdf"
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            # Create response with PDF data
            response = make_response(pdf_bytes)
            response.headers['Content-Type'] = 'application/pdf'
            # Use both filename and filename* for better browser compatibility
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{filename}'
            response.headers['Content-Length'] = str(len(pdf_bytes))
            response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
            response.headers['Pragma'] = 'no-cache'
            response.headers['Expires'] = '0'
            
            return response
            
        except Exception as e:
            import traceback
            error_msg = str(e)
            traceback.print_exc()
            # Log the error and return a proper error response
            app.logger.error(f"PDF export error: {error_msg}")
            flash(f"Error generating PDF: {error_msg}. Please check if reportlab is installed.", "danger")
            return redirect(url_for('view_assignments') if current_user.role == 'Student' else url_for('instructor_student_detail', student_id=current_user.id))

    @app.route('/export/csv')
    @login_required
    def export_csv():
        """Export student submissions to CSV"""
        from models.entities import Submission
        submissions = Submission.query.filter_by(student_id=current_user.id).order_by(Submission.created_at.asc()).all()
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Clean, readable headers in English
        writer.writerow(['Date', 'Submission Type', 'Score', 'Status', 'Feedback'])
        
        for sub in submissions:
            # Format score
            if sub.grade and sub.grade.score is not None:
                score = f"{sub.grade.score:.1f}"
                status = 'Graded' if sub.grade.instructor_approved else 'Pending'
            else:
                score = '-'
                status = 'Not Graded'
            
            # Format submission date to GMT+3 - more readable format
            if sub.created_at:
                sub_date_gmt3 = utc_to_gmt3(sub.created_at)
                date_str = sub_date_gmt3.strftime('%Y-%m-%d %H:%M') if sub_date_gmt3 else 'N/A'
            else:
                date_str = 'N/A'
            
            # Format submission type - readable English names
            type_map = {
                'WRITING': 'Writing',
                'SPEAKING': 'Speaking',
                'HANDWRITTEN': 'Handwritten',
                'QUIZ': 'Quiz'
            }
            submission_type = type_map.get(sub.submission_type, sub.submission_type.capitalize())
            
            # Format feedback - clean and concise
            if sub.grade and sub.grade.general_feedback:
                # Clean feedback: remove extra whitespace, limit length
                feedback = sub.grade.general_feedback.strip()
                # Limit to 150 characters for readability
                if len(feedback) > 150:
                    feedback = feedback[:147] + '...'
            else:
                feedback = '-'
            
            writer.writerow([
                date_str, 
                submission_type, 
                score, 
                status,
                feedback
            ])
        
        return Response(
            output.getvalue(),
            mimetype="text/csv",
            headers={"Content-disposition": "attachment; filename=academic_report.csv"}
        )

    @app.route('/export')
    @login_required
    def export_data():
        """Legacy export route for backwards compatibility"""
        format_type = request.args.get('format', 'csv')
        if format_type == 'pdf':
            return redirect(url_for('export_pdf'))
        else:
            return redirect(url_for('export_csv'))

    # ---  INSTRUCTOR DASHBOARD ---

    @app.route('/instructor/dashboard')
    @role_required('Instructor')
    def instructor_dashboard():
        from datetime import timedelta
        from collections import defaultdict
        from models.entities import Course, Enrollment
        
        # Get courses taught by this instructor
        instructor_courses = Course.query.filter_by(instructor_id=current_user.id, is_active=True).all()
        # Get student count for each course
        courses_with_students = []
        for course in instructor_courses:
            enrollments = Enrollment.query.filter_by(course_id=course.id, status='active').all()
            courses_with_students.append({
                'course': course,
                'student_count': len(enrollments)
            })
        
        all_subs = Submission.query.all()
        all_quizzes = Quiz.query.all()
        graded_subs = [s for s in all_subs if s.grade]
        class_avg = round(sum(s.grade.score for s in graded_subs) / len(graded_subs), 1) if graded_subs else 0.0
        active_count = len(set(s.student_id for s in all_subs))
        pending_count = len(all_subs) - len(graded_subs)
        
        # Calculate grade distribution
        grade_high = sum(1 for s in graded_subs if s.grade.score >= 75)
        grade_mid = sum(1 for s in graded_subs if 50 <= s.grade.score < 75)
        grade_low = sum(1 for s in graded_subs if s.grade.score < 50)
        
        # Prepare sparkline data for last 7 days
        today = datetime.utcnow().date()
        last_7_days = [today - timedelta(days=i) for i in range(6, -1, -1)]  # Last 7 days including today
        
        # Count submissions per day
        submissions_by_date = defaultdict(int)
        pending_by_date = defaultdict(int)
        avg_score_by_date = defaultdict(list)
        active_students_by_date = defaultdict(set)
        
        for sub in all_subs:
            sub_date = sub.created_at.date()
            if sub_date in last_7_days:
                submissions_by_date[sub_date] += 1
                if not sub.grade:
                    pending_by_date[sub_date] += 1
                if sub.grade:
                    avg_score_by_date[sub_date].append(sub.grade.score)
                active_students_by_date[sub_date].add(sub.student_id)
        
        # Create sparkline data arrays
        sparkline_data = {
            'submissions': [submissions_by_date.get(date, 0) for date in last_7_days],
            'pending': [pending_by_date.get(date, 0) for date in last_7_days],
            'class_avg': [round(sum(avg_score_by_date.get(date, [])) / len(avg_score_by_date.get(date, [])), 1) if avg_score_by_date.get(date, []) else 0.0 for date in last_7_days],
            'active_students': [len(active_students_by_date.get(date, set())) for date in last_7_days]
        }

        return render_template('instructor_dashboard.html', 
                               submissions=all_subs, 
                               quizzes=all_quizzes,
                               class_avg=class_avg, 
                               active_count=active_count,
                               pending_count=pending_count,
                               sparkline_data=sparkline_data,
                               grade_high=grade_high,
                               grade_mid=grade_mid,
                               grade_low=grade_low)

    @app.route('/instructor/courses')
    @role_required('Instructor')
    def instructor_courses():
        from models.entities import Course, Enrollment, LearningActivity, Submission
        # Get courses taught by this instructor
        instructor_courses_list = Course.query.filter_by(instructor_id=current_user.id, is_active=True).all()
        # Get stats for each course
        courses_with_students = []
        for course in instructor_courses_list:
            enrollments = Enrollment.query.filter_by(course_id=course.id, status='active').all()
            student_ids = [e.student_id for e in enrollments]
            
            # Count assignments for this course
            assignments = LearningActivity.query.filter_by(instructor_id=current_user.id).all()
            # Count assignments that have submissions from this course's students
            assignment_count = 0
            graded_count = 0
            total_scores = []
            
            if student_ids:
                # Get all submissions from students in this course
                course_submissions = Submission.query.filter(Submission.student_id.in_(student_ids)).all()
                
                # Count unique activities (assignments)
                activity_ids = set([s.activity_id for s in course_submissions if s.activity_id])
                assignment_count = len(activity_ids)
                
                # Count graded submissions
                graded_submissions = [s for s in course_submissions if s.grade and s.grade.score is not None]
                graded_count = len(graded_submissions)
                
                # Calculate average score
                for s in graded_submissions:
                    if s.grade and s.grade.score is not None:
                        total_scores.append(s.grade.score)
            
            avg_score = round(sum(total_scores) / len(total_scores), 1) if total_scores else 0.0
            
            courses_with_students.append({
                'course': course,
                'student_count': len(enrollments),
                'assignment_count': assignment_count,
                'graded_count': graded_count,
                'avg_score': avg_score
            })
        return render_template('instructor_courses.html', courses_with_students=courses_with_students)

    @app.route('/instructor/courses/<int:course_id>')
    @role_required('Instructor')
    def instructor_course_detail(course_id):
        from models.entities import Course, Enrollment, User, Submission, LearningActivity
        # Get course and verify instructor owns it
        course = Course.query.get_or_404(course_id)
        if course.instructor_id != current_user.id:
            flash('You do not have access to this course.', 'danger')
            return redirect(url_for('instructor_courses'))
        
        # Get all students enrolled in this course
        enrollments = Enrollment.query.filter_by(course_id=course_id, status='active').all()
        student_ids = [e.student_id for e in enrollments]
        students = User.query.filter(User.id.in_(student_ids), User.role == 'Student').all()
        
        # Get course statistics
        all_submissions = Submission.query.filter(Submission.student_id.in_(student_ids)).all() if student_ids else []
        graded_submissions = [s for s in all_submissions if s.grade and s.grade.score is not None]
        avg_score = round(sum(s.grade.score for s in graded_submissions) / len(graded_submissions), 1) if graded_submissions else 0.0
        
        # Get recent submissions (last 10, ordered by date)
        recent_submissions = Submission.query.filter(
            Submission.student_id.in_(student_ids)
        ).order_by(Submission.created_at.desc()).limit(10).all() if student_ids else []
        
        # Get assignments for this course (instructor's assignments)
        course_assignments = LearningActivity.query.filter_by(instructor_id=current_user.id).order_by(LearningActivity.created_at.desc()).limit(10).all()
        
        return render_template('instructor_course_detail.html', 
                             course=course, 
                             students=students,
                             student_count=len(students),
                             total_submissions=len(all_submissions),
                             graded_submissions=len(graded_submissions),
                             avg_score=avg_score,
                             recent_submissions=recent_submissions,
                             course_assignments=course_assignments)

    @app.route('/instructor/students')
    @role_required('Instructor')
    def instructor_students():
        students = User.query.filter_by(role='Student').all()
        return render_template('instructor_students.html', students=students)

    @app.route('/instructor/students/<int:student_id>')
    @role_required('Instructor')
    def instructor_student_detail(student_id):
        student = User.query.filter_by(id=student_id, role='Student').first_or_404()

        submissions = Submission.query.filter_by(student_id=student.id).order_by(Submission.created_at.desc()).all()
        quizzes = Quiz.query.filter_by(user_id=student.id).order_by(Quiz.id.desc()).all()
        goals = LearningGoal.query.filter_by(user_id=student.id).all()

        # Get student's enrolled courses
        from models.entities import Enrollment, Course
        enrollments = Enrollment.query.filter_by(student_id=student.id, status='active').all()
        student_courses = []
        for enrollment in enrollments:
            course = Course.query.get(enrollment.course_id)
            if course and course.is_active:
                student_courses.append(course)

        graded_subs = [s for s in submissions if s.grade]
        avg_score = round(sum(s.grade.score for s in graded_subs) / len(graded_subs), 1) if graded_subs else 0.0
        total_submissions = len(submissions)
        pending_submissions_list = [s for s in submissions if not s.grade]
        pending_submissions = len(pending_submissions_list)

        return render_template(
            'instructor_student_detail.html',
            student=student,
            submissions=submissions,
            quizzes=quizzes,
            goals=goals,
            student_courses=student_courses,
            avg_score=avg_score,
            total_submissions=total_submissions,
            pending_submissions=pending_submissions,
            pending_submissions_list=pending_submissions_list
        )

    @app.route('/instructor/export/pdf/<int:student_id>')
    @role_required('Instructor')
    def instructor_export_pdf(student_id):
        """Export student submissions to PDF for instructor"""
        # Verify student exists and is a student
        student = User.query.filter_by(id=student_id, role='Student').first_or_404()
        
        # Check if reportlab is available
        if not REPORTLAB_AVAILABLE:
            return "PDF generation requires reportlab library. Please install with: pip install reportlab", 500
        
        try:
            from models.entities import Submission
            
            # Get all submissions for the specified student, sorted chronologically (oldest first)
            submissions = db.session.query(Submission).filter_by(student_id=student_id).order_by(Submission.created_at.asc()).all()
            
            # Create PDF buffer
            buffer = io.BytesIO()
            try:
                p = canvas.Canvas(buffer, pagesize=letter)
                width, height = letter
                
                # Define consistent font sizes for entire document
                TITLE_FONT_SIZE = 18
                HEADER_INFO_FONT_SIZE = 11
                TABLE_HEADER_FONT_SIZE = 11
                TABLE_DATA_FONT_SIZE = 10
                
                # Header section
                p.setFont("Helvetica-Bold", TITLE_FONT_SIZE)
                p.drawString(100, height - 50, "AAFS AI - Academic Progress Report")
                p.setFont("Helvetica", HEADER_INFO_FONT_SIZE)
                p.drawString(100, height - 80, f"Student: {student.username}")
                # Use GMT+3 for generated time
                generated_time = get_gmt3_now()
                p.drawString(100, height - 100, f"Generated: {generated_time.strftime('%Y-%m-%d %H:%M')} (GMT+3)")
                p.line(100, height - 110, 500, height - 110)
                
                if submissions:
                    # Table headers - adjusted column positions
                    y = height - 130
                    p.setFont("Helvetica-Bold", TABLE_HEADER_FONT_SIZE)
                    p.drawString(100, y, "Submission Date")
                    p.drawString(280, y, "Type")
                    p.drawString(380, y, "Score")
                    p.line(100, y - 5, 500, y - 5)
                    
                    # Table rows
                    p.setFont("Helvetica", TABLE_DATA_FONT_SIZE)
                    y -= 25
                    
                    for sub in submissions:
                        # Check if we need a new page
                        if y < 100:
                            p.showPage()
                            # Re-draw headers on new page with same font sizes
                            y = height - 130
                            p.setFont("Helvetica-Bold", TABLE_HEADER_FONT_SIZE)
                            p.drawString(100, y, "Submission Date")
                            p.drawString(280, y, "Type")
                            p.drawString(380, y, "Score")
                            p.line(100, y - 5, 500, y - 5)
                            # Set data font for rows
                            p.setFont("Helvetica", TABLE_DATA_FONT_SIZE)
                            y -= 25
                        
                        try:
                            # Get score safely
                            if sub.grade and sub.grade.score is not None:
                                score = sub.grade.score
                            else:
                                score = 'N/A'
                            
                            # Format submission date with time in GMT+3
                            if sub.created_at:
                                sub_date_gmt3 = utc_to_gmt3(sub.created_at)
                                date_str = sub_date_gmt3.strftime('%Y-%m-%d %H:%M')
                            else:
                                date_str = 'N/A'
                            
                            # Format type safely
                            if sub.submission_type:
                                type_str = sub.submission_type.capitalize()
                            else:
                                type_str = 'Unknown'
                            
                            # Draw row
                            p.drawString(100, y, date_str)
                            p.drawString(280, y, type_str)
                            p.drawString(380, y, f"{score}%" if isinstance(score, (int, float)) else str(score))
                            
                            y -= 20
                        except Exception as e:
                            # Skip problematic submissions and continue
                            continue
                else:
                    # No submissions message
                    y = height - 130
                    p.setFont("Helvetica", HEADER_INFO_FONT_SIZE)
                    p.drawString(100, y, "No submissions found.")
                
                # Finalize and save PDF
                p.showPage()
                p.save()
            finally:
                # Ensure buffer is ready for reading
                buffer.seek(0)
            
            # Generate filename with GMT+3 date (safe ASCII only, use student ID instead of username)
            filename = f"academic_report_{student.id}_{get_gmt3_now().strftime('%Y%m%d')}.pdf"
            
            # Get PDF bytes
            pdf_bytes = buffer.getvalue()
            buffer.close()
            
            # Create response with PDF data
            response = make_response(pdf_bytes)
            response.headers['Content-Type'] = 'application/pdf'
            # Use both filename and filename* for better browser compatibility
            response.headers['Content-Disposition'] = f'attachment; filename="{filename}"; filename*=UTF-8\'\'{filename}'
            response.headers['Content-Length'] = str(len(pdf_bytes))
            response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
            response.headers['Pragma'] = 'no-cache'
            response.headers['Expires'] = '0'
            
            return response
            
        except Exception as e:
            import traceback
            error_msg = str(e)
            traceback.print_exc()
            # Log the error and return a proper error response
            app.logger.error(f"Instructor PDF export error: {error_msg}")
            flash(f"Error generating PDF: {error_msg}. Please check if reportlab is installed.", "danger")
            return redirect(url_for('instructor_student_detail', student_id=student_id))

    @app.route('/instructor/export/csv/<int:student_id>')
    @role_required('Instructor')
    def instructor_export_csv(student_id):
        """Export student submissions to CSV for instructor"""
        # Verify student exists and is a student
        student = User.query.filter_by(id=student_id, role='Student').first_or_404()
        
        from models.entities import Submission
        submissions = Submission.query.filter_by(student_id=student_id).order_by(Submission.created_at.asc()).all()
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Clean, readable headers in English
        writer.writerow(['Date', 'Submission Type', 'Score', 'Status', 'Feedback'])
        
        for sub in submissions:
            # Format score
            if sub.grade and sub.grade.score is not None:
                score = f"{sub.grade.score:.1f}"
                status = 'Graded' if sub.grade.instructor_approved else 'Pending'
            else:
                score = '-'
                status = 'Not Graded'
            
            # Format submission date to GMT+3 - more readable format
            if sub.created_at:
                sub_date_gmt3 = utc_to_gmt3(sub.created_at)
                date_str = sub_date_gmt3.strftime('%Y-%m-%d %H:%M') if sub_date_gmt3 else 'N/A'
            else:
                date_str = 'N/A'
            
            # Format submission type - readable English names
            type_map = {
                'WRITING': 'Writing',
                'SPEAKING': 'Speaking',
                'HANDWRITTEN': 'Handwritten',
                'QUIZ': 'Quiz'
            }
            submission_type = type_map.get(sub.submission_type, sub.submission_type.capitalize())
            
            # Format feedback - clean and concise
            if sub.grade and sub.grade.general_feedback:
                # Clean feedback: remove extra whitespace, limit length
                feedback = sub.grade.general_feedback.strip()
                # Limit to 150 characters for readability
                if len(feedback) > 150:
                    feedback = feedback[:147] + '...'
            else:
                feedback = '-'
            
            writer.writerow([
                date_str, 
                submission_type, 
                score, 
                status,
                feedback
            ])
        
        return Response(
            output.getvalue(),
            mimetype="text/csv",
            headers={"Content-disposition": f'attachment; filename=academic_report_{student.username}.csv'}
        )

    @app.route('/instructor/analytics')
    @role_required('Instructor')
    def instructor_analytics():
        from datetime import timedelta
        from collections import defaultdict
        
        students = User.query.filter_by(role='Student').all()
        all_subs = Submission.query.all()
        graded_subs = [s for s in all_subs if s.grade]
        
        total_students = len(students)
        total_submissions = len(all_subs)
        avg_score = round(sum(s.grade.score for s in graded_subs) / len(graded_subs), 1) if graded_subs else 0.0
        pending_reviews = len(all_subs) - len(graded_subs)
        
        # Chart data for last 7 days
        today = datetime.utcnow().date()
        last_7_days = [today - timedelta(days=i) for i in range(6, -1, -1)]
        submissions_by_date = defaultdict(int)
        
        for sub in all_subs:
            sub_date = sub.created_at.date()
            if sub_date in last_7_days:
                submissions_by_date[sub_date] += 1
        
        chart_labels = [date.strftime('%b %d') for date in last_7_days]
        submission_data = [submissions_by_date.get(date, 0) for date in last_7_days]
        
        # Top students
        top_students = sorted(students, key=lambda s: len(s.submissions) if s.submissions else 0, reverse=True)[:10]
        
        return render_template('instructor_analytics.html',
                               total_students=total_students,
                               total_submissions=total_submissions,
                               avg_score=avg_score,
                               pending_reviews=pending_reviews,
                               chart_labels=chart_labels,
                               submission_data=submission_data,
                               top_students=top_students)

    @app.route('/instructor/feedback')
    @role_required('Instructor')
    def instructor_feedback():
        student_id = request.args.get('student_id', type=int)
        filter_type = request.args.get('type', default=None, type=str)

        query = Submission.query.order_by(Submission.created_at.desc())

        if student_id:
            query = query.filter_by(student_id=student_id)

        if filter_type:
            # Expecting 'writing', 'speaking', 'handwritten' from UI
            submission_type = filter_type.upper()
            query = query.filter_by(submission_type=submission_type)

        submissions = query.all()

        return render_template(
            'instructor_feedback.html',
            submissions=submissions,
            selected_student_id=student_id,
            selected_type=filter_type or 'all'
        )

    @app.route('/instructor/submissions')
    @role_required('Instructor')
    def instructor_submissions():
        submissions = Submission.query.order_by(Submission.created_at.desc()).all()
        return render_template(
            'instructor_feedback.html',
            submissions=submissions,
            selected_student_id=None,
            selected_type='all'
        )

    @app.route('/instructor/pending')
    @role_required('Instructor')
    def instructor_pending():
        # Show submissions with AI grades that need instructor approval (instructor_approved=False)
        from models.entities import Grade
        submissions = Submission.query.join(Grade).filter(
            Grade.instructor_approved == False
        ).order_by(Submission.created_at.desc()).all()
        return render_template(
            'instructor_feedback.html',
            submissions=submissions,
            selected_student_id=None,
            selected_type='all'
        )
    
    @app.route('/instructor/assignments/<int:activity_id>')
    @role_required('Instructor')
    def instructor_assignment_detail(activity_id):
        activity = LearningActivity.query.get_or_404(activity_id)
        submissions = Submission.query.filter_by(activity_id=activity_id).order_by(Submission.created_at.desc()).all()
        
        total_submissions = len(submissions)
        # Graded = instructor approved
        graded_submissions = len([s for s in submissions if s.grade and s.grade.instructor_approved])
        # Pending = has AI grade but not approved yet
        pending_submissions = len([s for s in submissions if s.grade and not s.grade.instructor_approved])
        
        # Get students who submitted
        student_ids = set(s.student_id for s in submissions)
        students = User.query.filter(User.id.in_(student_ids)).all() if student_ids else []
        
        return render_template('instructor_assignment_detail.html',
                             activity=activity,
                             submissions=submissions,
                             total_submissions=total_submissions,
                             graded_submissions=graded_submissions,
                             pending_submissions=pending_submissions,
                             students=students)
    
    @app.route('/instructor/assignments/<int:activity_id>/edit', methods=['GET', 'POST'])
    @role_required('Instructor')
    def instructor_edit_assignment(activity_id):
        from models.entities import Course, User, Enrollment
        activity = LearningActivity.query.get_or_404(activity_id)
        
        if activity.instructor_id != current_user.id:
            flash("You don't have permission to edit this assignment.", "danger")
            return redirect(url_for('instructor_assignments'))
        
        # Get courses where this instructor teaches
        instructor_courses = Course.query.filter_by(instructor_id=current_user.id, is_active=True).all()
        course_ids = [c.id for c in instructor_courses]
        
        # Get students enrolled in these courses
        enrolled_students = []
        if course_ids:
            enrollments = Enrollment.query.filter(
                Enrollment.course_id.in_(course_ids),
                Enrollment.status == 'active'
            ).all()
            student_ids = list(set([e.student_id for e in enrollments]))
            enrolled_students = User.query.filter(
                User.id.in_(student_ids),
                User.role == 'Student'
            ).order_by(User.username.asc()).all()
        
        all_students = enrolled_students if enrolled_students else User.query.filter_by(role='Student').order_by(User.username.asc()).all()
        
        if request.method == 'POST':
            title = request.form.get('title')
            activity_type = request.form.get('activity_type')
            due_date_str = request.form.get('due_date')
            description = request.form.get('description')
            quiz_category = request.form.get('quiz_category') if activity_type == 'QUIZ' else None
            
            # Get assign_to setting
            assign_to = request.form.get('assign_to', 'all')
            student_id_str = request.form.get('student_id', '').strip()
            
            # Parse student_id
            student_id = None
            if assign_to == 'specific' and student_id_str:
                try:
                    student_id = int(student_id_str)
                    student = User.query.get(student_id)
                    if not student or student.role != 'Student':
                        flash('Invalid student selected.', 'danger')
                        return render_template('instructor_assignment_edit.html', activity=activity, courses=instructor_courses, students=all_students)
                except ValueError:
                    flash('Invalid student ID format.', 'danger')
                    return render_template('instructor_assignment_edit.html', activity=activity, courses=instructor_courses, students=all_students)
            
            # Get course IDs from form (multiple selection)
            course_ids_list = request.form.getlist('course_ids')
            course_ids_list = [int(cid) for cid in course_ids_list if cid and cid.strip()]
            
            # Validate that at least one course is selected
            if not course_ids_list:
                flash('Please select at least one course.', 'danger')
                return render_template('instructor_assignment_edit.html', activity=activity, courses=instructor_courses, students=all_students)
            
            # Verify all selected courses belong to this instructor
            valid_courses = Course.query.filter(
                Course.id.in_(course_ids_list),
                Course.instructor_id == current_user.id,
                Course.is_active == True
            ).all()
            if len(valid_courses) != len(course_ids_list):
                flash('Invalid course selection.', 'danger')
                return render_template('instructor_assignment_edit.html', activity=activity, courses=instructor_courses, students=all_students)
            
            if not title or not activity_type:
                flash('Title and type are required.', 'danger')
                return render_template('instructor_assignment_edit.html', activity=activity, courses=instructor_courses, students=all_students)
            
            activity.title = title
            activity.activity_type = activity_type
            activity.description = description
            activity.quiz_category = quiz_category
            activity.student_id = student_id  # Update student assignment
            
            due_date = None
            if due_date_str:
                try:
                    due_date = datetime.strptime(due_date_str, '%Y-%m-%d')
                except ValueError:
                    flash('Invalid date format. Use YYYY-MM-DD.', 'danger')
                    return render_template('instructor_assignment_edit.html', activity=activity, courses=instructor_courses, students=all_students)
            activity.due_date = due_date
            
            # Handle attachment file upload (update if new file is provided)
            if 'attachment' in request.files:
                attachment_file = request.files['attachment']
                if attachment_file and attachment_file.filename:
                    # Validate file size (max 10MB)
                    if attachment_file.content_length and attachment_file.content_length > 10 * 1024 * 1024:
                        flash('Attachment file size must be less than 10MB.', 'danger')
                        return render_template('instructor_assignment_edit.html', activity=activity, courses=instructor_courses, students=all_students)
                    
                    # Delete old attachment if exists
                    if activity.attachment_path:
                        old_file_path = os.path.join(app.config['UPLOAD_FOLDER'], activity.attachment_path)
                        if os.path.exists(old_file_path):
                            try:
                                os.remove(old_file_path)
                            except Exception as e:
                                app.logger.warning(f"Could not delete old attachment: {e}")
                    
                    # Save new attachment file
                    filename = secure_filename(attachment_file.filename)
                    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
                    filename = f"{timestamp}_{filename}"
                    attachment_dir = os.path.join(app.config['UPLOAD_FOLDER'], 'assignments')
                    os.makedirs(attachment_dir, exist_ok=True)
                    file_path = os.path.join(attachment_dir, filename)
                    attachment_file.save(file_path)
                    # Use forward slash for web URLs (works on all platforms)
                    activity.attachment_path = 'assignments/' + filename
                    activity.attachment_filename = attachment_file.filename
            
            # Update course assignments
            from services.activity_service import ActivityService
            ActivityService.update_activity_courses(activity_id, course_ids_list)
            
            db.session.commit()
            flash('Assignment updated successfully!', 'success')
            return redirect(url_for('instructor_assignment_detail', activity_id=activity_id))
        
        return render_template('instructor_assignment_edit.html', activity=activity, courses=instructor_courses, students=all_students)
    
    @app.route('/instructor/assignments/<int:activity_id>/delete', methods=['POST'])
    @role_required('Instructor')
    def instructor_delete_assignment(activity_id):
        activity = LearningActivity.query.get_or_404(activity_id)
        
        if activity.instructor_id != current_user.id:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'message': 'Permission denied'}), 403
            flash("You don't have permission to delete this assignment.", "danger")
            return redirect(url_for('instructor_assignments'))
        
        # Delete associated submissions and grades
        submissions = Submission.query.filter_by(activity_id=activity_id).all()
        for submission in submissions:
            if submission.grade:
                db.session.delete(submission.grade)
            db.session.delete(submission)
        
        db.session.delete(activity)
        db.session.commit()
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return jsonify({'success': True, 'message': 'Assignment deleted successfully!'}), 200
        
        flash('Assignment deleted successfully!', 'success')
        return redirect(url_for('instructor_assignments'))
    
    @app.route('/instructor/submissions/<int:submission_id>/approve', methods=['POST'])
    @role_required('Instructor')
    def approve_submission(submission_id):
        submission = Submission.query.get_or_404(submission_id)
        
        if not submission.grade:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return jsonify({'success': False, 'message': 'No grade found for this submission'}), 400
            flash('No grade found for this submission.', 'danger')
            return redirect(url_for('instructor_pending'))
        
        success = GradingService.approve_grade(submission_id)
        
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            if success:
                return jsonify({'success': True, 'message': 'Grade approved successfully!'}), 200
            else:
                return jsonify({'success': False, 'message': 'Failed to approve grade'}), 400
        
        if success:
            flash('Grade approved successfully!', 'success')
        else:
            flash('Failed to approve grade.', 'danger')
        
        return redirect(url_for('instructor_pending'))

    @app.route('/instructor/questions', methods=['GET', 'POST'])
    @role_required('Instructor')
    def manage_questions():
        if request.method == 'POST':
            question_text = request.form.get('question_text')
            option_a = request.form.get('option_a')
            option_b = request.form.get('option_b')
            option_c = request.form.get('option_c', '')
            option_d = request.form.get('option_d', '')
            correct_answer = request.form.get('correct_answer')
            category = request.form.get('category', 'grammar')
            
            if question_text and option_a and option_b and correct_answer:
                new_question = Question(
                    question_text=question_text,
                    option_a=option_a,
                    option_b=option_b,
                    option_c=option_c if option_c else None,
                    option_d=option_d if option_d else None,
                    correct_answer=correct_answer.upper(),
                    category=category
                )
                db.session.add(new_question)
                db.session.commit()
                flash("Question added successfully!", "success")
                return redirect(url_for('manage_questions'))
            else:
                flash("Please fill in all required fields.", "danger")
        
        questions = Question.query.all()
        return render_template('manage_questions.html', questions=questions)

    # --- SUBMISSION ROUTES ---

    @app.route('/submit/writing', methods=['GET', 'POST'])
    @role_required('Student')
    def submit_writing():
        activity_id = request.args.get('activity_id') or request.form.get('activity_id')
        if request.method == 'POST':
            text_content = request.form.get('text_content', '').strip()
            file = request.files.get('file')
            
            # Handle file upload using SubmissionService
            if file and file.filename != '':
                # Validate file format
                if not SubmissionService.validate_file_format(file.filename):
                    flash("Invalid file format. Please upload .docx, .pdf, .txt, or image files.", "danger")
                    return redirect(url_for('submit_writing'))
                
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                
                if filename.endswith('.docx'):
                    doc = docx.Document(file_path)
                    text_content = "\n".join([p.text for p in doc.paragraphs])
                else:
                    try:
                        with open(file_path, 'r', encoding='utf-8') as f: 
                            text_content = f.read()
                    except:
                        with open(file_path, 'r', encoding='latin-1') as f: 
                            text_content = f.read()
            
            # If no text content but file is provided, extract text from file
            if not text_content and file and file.filename != '':
                # File processing is already done above, text_content should be set
                # But if it's still empty, try to extract again
                if not text_content:
                    filename = secure_filename(file.filename)
                    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                    if os.path.exists(file_path):
                        if filename.endswith('.docx'):
                            doc = docx.Document(file_path)
                            text_content = "\n".join([p.text for p in doc.paragraphs])
                        else:
                            try:
                                with open(file_path, 'r', encoding='utf-8') as f:
                                    text_content = f.read()
                            except:
                                with open(file_path, 'r', encoding='latin-1') as f:
                                    text_content = f.read()
            
            # Check if we have text content (either from input or file)
            if not text_content:
                flash("Please provide some text or upload a file to analyze.", "danger")
                return redirect(url_for('submit_writing'))
            
            # Create submission using SubmissionService
            new_sub, error_msg = SubmissionService.save_submission_text(
                student_id=current_user.id,
                activity_id=activity_id,
                submission_type='WRITING',
                text_content=text_content,
                file_path=file.filename if file else None
            )

            if not new_sub:
                flash(error_msg or "Failed to create submission.", "danger")
                return render_template('submit_writing.html', submitted_text=text_content)

            # Check if AI is enabled
            if not AIService._is_ai_enabled():
                flash("AI features are currently disabled by administrator.", "danger")
                return render_template('submit_writing.html', submitted_text=text_content)
            
            # Analyze with AI
            print(f"Starting AI analysis for submission {new_sub.id}")
            ai_res = AIService.evaluate_writing(text_content)
            
            # Process evaluation using GradingService
            if ai_res and ai_res.get('score') is not None:
                success = GradingService.process_evaluation(new_sub.id, ai_res)
                if success:
                    # Update goal progress for Writing category
                    GoalService.update_goal_progress(current_user.id, 'Writing')
                    NotificationService.notify_grade_ready(current_user.id, new_sub.id)
                    flash("Submission analyzed successfully!", "success")
                    # Reload the submission with grade to show results on the same page
                    new_sub = Submission.query.get(new_sub.id)
                    return render_template('submit_writing.html', 
                                         grade=new_sub.grade,
                                         submitted_text=text_content,
                                         analysis_results=ai_res,
                                         submission_id=new_sub.id,
                                         activity_id=activity_id)
                else:
                    flash("Failed to save grade.", "danger")
            else:
                error_msg = ai_res.get('general_feedback', 'Unknown error') if ai_res else 'No response from AI'
                flash(f"Analysis failed: {error_msg}", "danger")
            
            return render_template('submit_writing.html', submitted_text=text_content)
        
        # GET request - check if there's a submission_id to show results
        submission_id = request.args.get('submission_id', type=int)
        grade = None
        submitted_text = None
        submission_activity_id = None
        
        if submission_id:
            submission = Submission.query.filter_by(id=submission_id, student_id=current_user.id).first()
            if submission:
                grade = submission.grade
                submitted_text = submission.text_content
                submission_activity_id = submission.activity_id
        
        return render_template('submit_writing.html', 
                              grade=grade, 
                              submitted_text=submitted_text,
                              submission_id=submission_id,
                              activity_id=submission_activity_id or activity_id)

    @app.route('/submit/writing/<int:submission_id>/finalize', methods=['POST'])
    @role_required('Student')
    def finalize_writing_submission(submission_id):
        """Finalize a writing submission after analysis"""
        submission = Submission.query.filter_by(
            id=submission_id,
            student_id=current_user.id
        ).first()
        
        if not submission:
            flash("Submission not found.", "danger")
            return redirect(url_for('submit_writing'))
        
        # Mark submission as completed
        submission.status = 'COMPLETED'
        db.session.commit()
        
        flash("Assignment submitted successfully! ✅", "success")
        return redirect(url_for('view_feedback', submission_id=submission_id))

    @app.route('/submit/handwritten', methods=['GET', 'POST'])
    @role_required('Student')
    def submit_handwritten():
        activity_id = request.args.get('activity_id')
        image_path = None
        extracted_text = None
        uploaded_filename = None
        grade = None
        error_message = None
        
        if request.method == 'POST':
            file = request.files.get('file')
            if file and file.filename != '':
                file_ext = os.path.splitext(file.filename)[1].lower()
                allowed_extensions = {'.pdf', '.jpg', '.jpeg', '.png', '.gif'}
                if file_ext not in allowed_extensions:
                    flash("Invalid file format. Please upload PDF or image files.", "danger")
                    return render_template('submit_handwritten.html', 
                                         image_path=None,
                                         extracted_text=None,
                                         uploaded_filename=None,
                                         grade=None,
                                         error_message="Invalid file format. Please upload PDF or image files.")

                if file_ext == '.pdf':
                    try:
                        import fitz  # noqa: F401
                    except Exception:
                        error_message = "PDF OCR requires PyMuPDF. Install it and restart the server."
                        return render_template('submit_handwritten.html', 
                                             image_path=None,
                                             extracted_text=None,
                                             uploaded_filename=file.filename,
                                             grade=None,
                                             error_message=error_message)

                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
                uploaded_filename = filename
                
                extracted_text = OCRService.extract_text_from_image(file_path)
                if not extracted_text:
                    flash("Failed to extract text from image. Please upload a clearer image with better handwriting.", "danger")
                    return render_template('submit_handwritten.html', 
                                         image_path=None,
                                         extracted_text=None,
                                         uploaded_filename=uploaded_filename,
                                         grade=None,
                                         error_message="No text could be extracted. Check Tesseract setup or try a clearer PDF/image.")
                
                if extracted_text:
                    # Save submission using SubmissionService
                    new_sub, error_msg = SubmissionService.save_submission_text(
                        student_id=current_user.id,
                        activity_id=activity_id,
                        submission_type='HANDWRITTEN',
                        text_content=extracted_text,
                        file_path=filename
                    )
                    
                    if not new_sub:
                        flash(error_msg or "Failed to create submission.", "danger")
                        return render_template('submit_handwritten.html', 
                                             image_path=None,
                                             extracted_text=None)
                    
                    # Check if AI is enabled
                    if AIService._is_ai_enabled():
                        ai_res = AIService.evaluate_writing(extracted_text)
                        if ai_res:
                            # Process evaluation using GradingService
                            success = GradingService.process_evaluation(new_sub.id, ai_res)
                            if success:
                                NotificationService.notify_grade_ready(current_user.id, new_sub.id)
                                new_sub = Submission.query.get(new_sub.id)
                                grade = new_sub.grade
                    else:
                        flash("AI features are currently disabled by administrator.", "danger")
                    
                    # Set image path for display (relative to static folder)
                    if file_ext in {'.jpg', '.jpeg', '.png', '.gif'}:
                        image_path = f"uploads/{filename}"
                    flash("Image processed successfully!", "success")
                    
        return render_template('submit_handwritten.html', 
                               image_path=image_path,
                               extracted_text=extracted_text,
                               uploaded_filename=uploaded_filename,
                               grade=grade,
                               error_message=error_message)

    @app.route('/history')
    @login_required
    def history():
        filter_type = request.args.get('filter')

        # For quiz-only history, use Quiz table
        if filter_type == 'quiz':
            quizzes = QuizRepository.get_quizzes(user_id=current_user.id)
            submissions = []
        else:
            query = Submission.query.filter_by(student_id=current_user.id)
            
            if filter_type:
                if filter_type == 'speaking':
                    query = query.filter_by(submission_type='SPEAKING')
                elif filter_type == 'writing':
                    query = query.filter_by(submission_type='WRITING')
                elif filter_type == 'handwritten':
                    query = query.filter_by(submission_type='HANDWRITTEN')
            
            submissions = query.order_by(Submission.created_at.desc()).all()

            # In "All" view also include quizzes; for filtered speaking/writing/handwritten, hide them
            if not filter_type or filter_type == 'all':
                quizzes = QuizRepository.get_quizzes(user_id=current_user.id)
            else:
                quizzes = []
        
        # For assignments history
        assignments = []
        activity_submissions = {}
        if filter_type == 'assignments':
            # Get all completed assignments (submitted ones)
            user_subs_all = Submission.query.filter_by(student_id=current_user.id).all()
            completed_activity_ids = set(s.activity_id for s in user_subs_all if s.activity_id)
            
            # Get activities that have been submitted
            assignments = LearningActivity.query.filter(
                LearningActivity.id.in_(completed_activity_ids)
            ).order_by(LearningActivity.due_date.desc()).all()
            
            # Get submissions for these activities
            for sub in user_subs_all:
                if sub.activity_id:
                    if sub.activity_id not in activity_submissions:
                        activity_submissions[sub.activity_id] = []
                    activity_submissions[sub.activity_id].append(sub)
        
        return render_template('history.html', 
                             submissions=submissions, 
                             quizzes=quizzes,
                             assignments=assignments if filter_type == 'assignments' else [],
                             activity_submissions=activity_submissions if filter_type == 'assignments' else {},
                             now=datetime.utcnow())

    @app.route('/feedback/<int:submission_id>')
    @login_required
    def view_feedback(submission_id):
        # Use FeedbackRepository to find feedback
        grade = FeedbackRepository.find_feedback_by_submission_id(submission_id)
        sub = Submission.query.get_or_404(submission_id)
        
        # Ensure user can only view their own submissions (unless instructor)
        if current_user.role != 'Instructor' and sub.student_id != current_user.id:
            flash("You don't have permission to view this report.", "error")
            return redirect(url_for('dashboard'))
        
        # Load activity if submission is part of an assignment
        if sub.activity_id:
            sub.activity = LearningActivity.query.get(sub.activity_id)
        
        return render_template('feedback.html', submission=sub)

    @app.route('/instructor/adjust-grade/<int:submission_id>', methods=['GET', 'POST'])
    @login_required
    def adjust_grade(submission_id):
        from models.entities import Submission, Grade
        submission = db.session.get(Submission, submission_id)
        if not submission:
            flash('Submission not found', 'danger')
            return redirect(url_for('instructor_dashboard'))
        
        if not submission.grade:
            new_grade = Grade(submission_id=submission.id, score=0.0)
            db.session.add(new_grade)
            db.session.commit()
            submission = db.session.get(Submission, submission_id)  # Refresh to get the new grade
        
        if request.method == 'POST':
            try:
                score = request.form.get('score', type=float)
                feedback = request.form.get('feedback', '')
                
                if score is not None and 0 <= score <= 100:
                    submission.grade.score = score
                    submission.grade.general_feedback = feedback
                    submission.grade.instructor_approved = True
                    db.session.commit()
                    flash('Success: Evaluation updated!', 'success')
                    return redirect(url_for('instructor_student_detail', student_id=submission.student_id))
                else:
                    flash('Invalid score. Please enter a value between 0 and 100.', 'danger')
            except Exception as e:
                db.session.rollback()
                flash(f'Error: {str(e)}', 'danger')
        
        return render_template('adjust_grade.html', submission=submission)

    @app.route('/delete_submission/<int:submission_id>', methods=['POST'])
    @login_required
    def delete_submission(submission_id):
        from flask import jsonify
        sub = Submission.query.get_or_404(submission_id)
        # Ensure user can only delete their own submissions
        if sub.student_id != current_user.id:
            return jsonify({'success': False, 'error': 'Permission denied'}), 403
        
        try:
            # Delete associated grade if exists
            if sub.grade:
                db.session.delete(sub.grade)
            
            # Delete submission
            db.session.delete(sub)
            db.session.commit()
            
            return jsonify({'success': True})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/delete_quiz/<int:quiz_id>', methods=['POST'])
    @login_required
    def delete_quiz(quiz_id):
        from flask import jsonify
        quiz = Quiz.query.get_or_404(quiz_id)
        # Ensure user can only delete their own quizzes
        if quiz.user_id != current_user.id:
            return jsonify({'success': False, 'error': 'Permission denied'}), 403
        
        try:
            # Delete associated quiz details if exists
            from models.entities import QuizDetail
            quiz_details = QuizDetail.query.filter_by(quiz_id=quiz.id).all()
            for detail in quiz_details:
                db.session.delete(detail)
            
            # Delete quiz
            db.session.delete(quiz)
            db.session.commit()
            
            return jsonify({'success': True})
        except Exception as e:
            db.session.rollback()
            return jsonify({'success': False, 'error': str(e)}), 500

    @app.route('/privacy')
    def privacy():
        return render_template('privacy.html')

    @app.route('/terms')
    def terms():
        return render_template('terms.html')

    # --- ADMIN ROUTES ---
    
    @app.route('/admin/dashboard')
    @role_required('Admin')
    def admin_dashboard():
        stats = AdminService.get_user_statistics()
        recent_users = AdminRepository.get_all_users()[:10]
        recent_courses = AdminRepository.get_all_courses()[:5]
        return render_template('admin_dashboard.html', stats=stats, recent_users=recent_users, recent_courses=recent_courses)
    
    @app.route('/admin/users')
    @role_required('Admin')
    def admin_users():
        users = AdminRepository.get_all_users()
        role_filter = request.args.get('role', 'all')
        if role_filter != 'all':
            users = [u for u in users if u.role == role_filter]
        return render_template('admin_users.html', users=users, role_filter=role_filter)
    
    @app.route('/admin/users/create', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_create_user():
        if request.method == 'POST':
            username = request.form.get('username')
            email = request.form.get('email')
            password = request.form.get('password')
            role = request.form.get('role', 'Student')
            
            errors = AdminService.validate_user_data(username, email, password, role)
            if errors:
                for error in errors:
                    flash(error, 'danger')
                return render_template('admin_user_create.html')
            
            try:
                AdminRepository.create_user(username, email, password, role)
                flash('User created successfully!', 'success')
                return redirect(url_for('admin_users'))
            except Exception as e:
                flash(f'Error creating user: {str(e)}', 'danger')
        
        return render_template('admin_user_create.html')
    
    @app.route('/admin/users/<int:user_id>/edit', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_edit_user(user_id):
        user = AdminRepository.get_user_by_id(user_id)
        if not user:
            flash('User not found', 'danger')
            return redirect(url_for('admin_users'))
        
        if request.method == 'POST':
            username = request.form.get('username', '').strip()
            email = request.form.get('email', '').strip()
            password = request.form.get('password', '').strip()  # Optional
            role = request.form.get('role', '').strip()
            
            # Validate role is provided
            if not role:
                flash('Role is required', 'danger')
                return render_template('admin_user_edit.html', user=user)
            
            # Manual validation (don't use AdminService.validate_user_data for updates as it does duplicate checks without excluding current user)
            errors = []
            
            if not username or len(username) == 0:
                errors.append("Username is required")
            elif len(username) > 50:
                errors.append("Username must be 50 characters or less")
            
            if not email or len(email) == 0:
                errors.append("Email is required")
            elif '@' not in email:
                errors.append("Invalid email format")
            
            if password and len(password) < 6:
                errors.append("Password must be at least 6 characters")
            
            if role and role not in ['Student', 'Instructor', 'Admin']:
                errors.append("Invalid role. Must be Student, Instructor, or Admin")
            
            # Check if username/email conflicts with OTHER users (exclude current user)
            existing_username = User.query.filter_by(username=username).first()
            if existing_username and existing_username.id != user_id:
                errors.append("Username already exists")
            
            existing_email = User.query.filter_by(email=email).first()
            if existing_email and existing_email.id != user_id:
                errors.append("Email already exists")
            
            if errors:
                for error in errors:
                    flash(error, 'danger')
                return render_template('admin_user_edit.html', user=user)
            
            try:
                # Verify user_id is valid and different from current_user to prevent accidental self-update
                if user_id == current_user.id:
                    flash('You cannot edit your own account from this page. Use Settings page instead.', 'warning')
                    return redirect(url_for('admin_users'))
                
                # Always update role if provided (even if same value)
                # Ensure we're updating the correct user by explicitly passing user_id
                updated_user = AdminRepository.update_user(user_id, username, email, password if password else None, role)
                if updated_user:
                    # Verify the updated user is the one we intended to update
                    if updated_user.id != user_id:
                        db.session.rollback()
                        flash('Error: Wrong user was updated. Please try again.', 'danger')
                        return redirect(url_for('admin_users'))
                    flash(f'User {updated_user.username} updated successfully! Role changed to {role}.', 'success')
                else:
                    flash('User not found', 'danger')
                return redirect(url_for('admin_users'))
            except Exception as e:
                db.session.rollback()
                flash(f'Error updating user: {str(e)}', 'danger')
                import traceback
                traceback.print_exc()
        
        return render_template('admin_user_edit.html', user=user)
    
    @app.route('/admin/users/<int:user_id>/delete', methods=['POST'])
    @role_required('Admin')
    def admin_delete_user(user_id):
        if user_id == current_user.id:
            flash('You cannot delete your own account', 'danger')
            return redirect(url_for('admin_users'))
        
        try:
            result = AdminRepository.delete_user(user_id)
            if result:
                flash('User deleted successfully!', 'success')
            else:
                flash('User not found or could not be deleted.', 'danger')
        except Exception as e:
            db.session.rollback()
            flash(f'Error deleting user: {str(e)}', 'danger')
        
        return redirect(url_for('admin_users'))
    
    @app.route('/admin/courses')
    @role_required('Admin')
    def admin_courses():
        courses = AdminRepository.get_all_courses()
        instructors = AdminRepository.get_users_by_role('Instructor')
        return render_template('admin_courses.html', courses=courses, instructors=instructors)
    
    @app.route('/admin/courses/create', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_create_course():
        instructors = AdminRepository.get_users_by_role('Instructor')
        students = AdminRepository.get_users_by_role('Student')
        
        if request.method == 'POST':
            name = request.form.get('name')
            code = request.form.get('code')
            description = request.form.get('description')
            instructor_id = request.form.get('instructor_id', type=int) or None
            is_active = request.form.get('is_active') == 'on'
            # Get selected student IDs (multiple selection)
            student_ids = request.form.getlist('student_ids')
            
            errors = AdminService.validate_course_data(name, code, description)
            if errors:
                for error in errors:
                    flash(error, 'danger')
                return render_template('admin_course_create.html', instructors=instructors, students=students)
            
            try:
                # Create the course
                course = AdminRepository.create_course(name, code, description, instructor_id, is_active)
                
                # Enroll selected students
                enrolled_count = 0
                if student_ids:
                    for student_id_str in student_ids:
                        try:
                            student_id = int(student_id_str)
                            enrollment = AdminRepository.create_enrollment(student_id, course.id, 'active')
                            if enrollment:
                                enrolled_count += 1
                        except (ValueError, Exception) as e:
                            # Skip invalid student IDs, continue with others
                            continue
                
                if enrolled_count > 0:
                    flash(f'Course created successfully! {enrolled_count} student(s) enrolled.', 'success')
                else:
                    flash('Course created successfully!', 'success')
                return redirect(url_for('admin_courses'))
            except Exception as e:
                flash(f'Error creating course: {str(e)}', 'danger')
                import traceback
                traceback.print_exc()
        
        return render_template('admin_course_create.html', instructors=instructors, students=students)
    
    @app.route('/admin/courses/<int:course_id>/edit', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_edit_course(course_id):
        course = AdminRepository.get_course_by_id(course_id)
        if not course:
            flash('Course not found', 'danger')
            return redirect(url_for('admin_courses'))
        
        instructors = AdminRepository.get_users_by_role('Instructor')
        students = AdminRepository.get_users_by_role('Student')
        # Get already enrolled student IDs
        existing_enrollments = AdminRepository.get_enrollments_by_course(course_id)
        enrolled_student_ids = {e.student_id for e in existing_enrollments}
        # Filter out already enrolled students
        available_students = [s for s in students if s.id not in enrolled_student_ids]
        
        if request.method == 'POST':
            name = request.form.get('name', '').strip()
            code = request.form.get('code', '').strip()
            description = request.form.get('description', '').strip()
            instructor_id = request.form.get('instructor_id', type=int) or None
            is_active = request.form.get('is_active') == 'on'
            # Get selected student IDs (for adding new enrollments)
            student_ids = request.form.getlist('student_ids')
            
            # Manual validation for edit (exclude current course from duplicate checks)
            errors = []
            
            if not name or len(name) == 0:
                errors.append("Course name is required")
            
            if not code or len(code) == 0:
                errors.append("Course code is required")
            
            # Check if code conflicts with OTHER courses (exclude current course)
            existing_course = Course.query.filter_by(code=code).first()
            if existing_course and existing_course.id != course_id:
                errors.append("Course code already exists")
            
            if errors:
                for error in errors:
                    flash(error, 'danger')
                return render_template('admin_course_edit.html', course=course, instructors=instructors, 
                                     available_students=available_students, enrolled_student_ids=enrolled_student_ids)
            
            try:
                # Update the course - ensure all fields are passed explicitly
                # Note: description can be empty string, which is valid
                updated_course = AdminRepository.update_course(
                    course_id, 
                    name=name, 
                    code=code, 
                    description=description,  # Can be empty string
                    instructor_id=instructor_id, 
                    is_active=is_active
                )
                
                if not updated_course:
                    db.session.rollback()
                    flash('Course not found', 'danger')
                    return redirect(url_for('admin_courses'))
                
                # Verify the updated course is the one we intended to update
                if updated_course.id != course_id:
                    db.session.rollback()
                    flash('Error: Wrong course was updated. Please try again.', 'danger')
                    return redirect(url_for('admin_courses'))
                
                # Enroll selected students (new enrollments only)
                enrolled_count = 0
                if student_ids:
                    for student_id_str in student_ids:
                        try:
                            student_id = int(student_id_str)
                            # Only enroll if not already enrolled
                            if student_id not in enrolled_student_ids:
                                enrollment = AdminRepository.create_enrollment(student_id, course_id, 'active')
                                if enrollment:
                                    enrolled_count += 1
                        except (ValueError, Exception) as e:
                            # Skip invalid student IDs, continue with others
                            continue
                
                if enrolled_count > 0:
                    flash(f'Course "{updated_course.name}" updated successfully! {enrolled_count} new student(s) enrolled.', 'success')
                else:
                    flash(f'Course "{updated_course.name}" updated successfully!', 'success')
                return redirect(url_for('admin_courses'))
            except Exception as e:
                db.session.rollback()
                flash(f'Error updating course: {str(e)}', 'danger')
                import traceback
                traceback.print_exc()
        
        return render_template('admin_course_edit.html', course=course, instructors=instructors, 
                             available_students=available_students, enrolled_student_ids=enrolled_student_ids)
    
    @app.route('/admin/courses/<int:course_id>/delete', methods=['POST'])
    @role_required('Admin')
    def admin_delete_course(course_id):
        try:
            AdminRepository.delete_course(course_id)
            flash('Course deleted successfully!', 'success')
        except Exception as e:
            flash(f'Error deleting course: {str(e)}', 'danger')
        
        return redirect(url_for('admin_courses'))
    
    @app.route('/admin/enrollments')
    @role_required('Admin')
    def admin_enrollments():
        enrollments = AdminRepository.get_all_enrollments()
        course_filter = request.args.get('course_id', type=int)
        student_filter = request.args.get('student_id', type=int)
        
        if course_filter:
            enrollments = [e for e in enrollments if e.course_id == course_filter]
        if student_filter:
            enrollments = [e for e in enrollments if e.student_id == student_filter]
        
        courses = AdminRepository.get_all_courses()
        students = AdminRepository.get_users_by_role('Student')
        
        return render_template('admin_enrollments.html', 
                             enrollments=enrollments, 
                             courses=courses, 
                             students=students,
                             course_filter=course_filter,
                             student_filter=student_filter)
    
    @app.route('/admin/enrollments/create', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_create_enrollment():
        courses = AdminRepository.get_all_courses()
        students = AdminRepository.get_users_by_role('Student')
        
        if request.method == 'POST':
            # Get multiple student IDs from form
            student_ids = request.form.getlist('student_ids')
            course_id = request.form.get('course_id', type=int)
            status = request.form.get('status', 'active')
            
            if not student_ids or not course_id:
                flash('At least one student and a course are required', 'danger')
                return render_template('admin_enrollment_create.html', courses=courses, students=students)
            
            # Convert string IDs to integers
            try:
                student_ids = [int(sid) for sid in student_ids if sid]
            except (ValueError, TypeError):
                flash('Invalid student selection', 'danger')
                return render_template('admin_enrollment_create.html', courses=courses, students=students)
            
            if not student_ids:
                flash('Please select at least one student', 'danger')
                return render_template('admin_enrollment_create.html', courses=courses, students=students)
            
            # Create enrollments for each student
            success_count = 0
            failed_students = []
            
            for student_id in student_ids:
                try:
                    result = AdminRepository.create_enrollment(student_id, course_id, status)
                    if result:
                        success_count += 1
                    else:
                        # Get student username for error message
                        student = User.query.get(student_id)
                        student_name = student.username if student else f"Student ID {student_id}"
                        failed_students.append(student_name)
                except Exception as e:
                    student = User.query.get(student_id)
                    student_name = student.username if student else f"Student ID {student_id}"
                    failed_students.append(f"{student_name} (Error: {str(e)})")
            
            # Show appropriate message based on results
            if success_count > 0 and len(failed_students) == 0:
                if success_count == 1:
                    flash('Enrollment created successfully!', 'success')
                else:
                    flash(f'{success_count} enrollments created successfully!', 'success')
                return redirect(url_for('admin_enrollments'))
            elif success_count > 0 and len(failed_students) > 0:
                # Partial success
                error_msg = f'{success_count} enrollment(s) created. Failed for: {", ".join(failed_students)}'
                flash(error_msg, 'warning')
                return render_template('admin_enrollment_create.html', courses=courses, students=students)
            else:
                # All failed
                if len(failed_students) == 1 and 'already enrolled' in str(failed_students[0]).lower():
                    flash('Student is already enrolled in this course', 'danger')
                else:
                    flash(f'Failed to create enrollments: {", ".join(failed_students)}', 'danger')
                return render_template('admin_enrollment_create.html', courses=courses, students=students)
        
        return render_template('admin_enrollment_create.html', courses=courses, students=students)
    
    @app.route('/admin/enrollments/<int:enrollment_id>/edit', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_edit_enrollment(enrollment_id):
        enrollment = AdminRepository.get_enrollment_by_id(enrollment_id)
        if not enrollment:
            flash('Enrollment not found', 'danger')
            return redirect(url_for('admin_enrollments'))
        
        if request.method == 'POST':
            status = request.form.get('status')
            
            try:
                AdminRepository.update_enrollment(enrollment_id, status)
                flash('Enrollment updated successfully!', 'success')
                return redirect(url_for('admin_enrollments'))
            except Exception as e:
                flash(f'Error updating enrollment: {str(e)}', 'danger')
        
        return render_template('admin_enrollment_edit.html', enrollment=enrollment)
    
    @app.route('/admin/enrollments/<int:enrollment_id>/delete', methods=['POST'])
    @role_required('Admin')
    def admin_delete_enrollment(enrollment_id):
        try:
            AdminRepository.delete_enrollment(enrollment_id)
            flash('Enrollment deleted successfully!', 'success')
        except Exception as e:
            flash(f'Error deleting enrollment: {str(e)}', 'danger')
        
        return redirect(url_for('admin_enrollments'))
    
    @app.route('/admin/settings')
    @role_required('Admin')
    def admin_settings():
        # Redirect to dashboard - Settings page has been removed
        return redirect(url_for('admin_dashboard'))
    
    @app.route('/admin/settings/update', methods=['POST'])
    @role_required('Admin')
    def admin_update_setting():
        # Redirect to dashboard - Settings page has been removed
        return redirect(url_for('admin_dashboard'))
    
    @app.route('/admin/ai-integrations')
    @role_required('Admin')
    def admin_ai_integrations():
        integrations = AdminRepository.get_all_ai_integrations()
        # Check if GEMINI_API_KEY is configured in environment
        gemini_api_key_configured = bool(os.getenv('GEMINI_API_KEY'))
        return render_template('admin_ai_integrations.html', 
                             integrations=integrations,
                             gemini_api_key_configured=gemini_api_key_configured)
    
    @app.route('/admin/ai-integrations/create', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_create_ai_integration():
        if request.method == 'POST':
            integration_name = request.form.get('integration_name')
            api_key = request.form.get('api_key')
            is_active = request.form.get('is_active') == 'on'
            api_endpoint = request.form.get('api_endpoint')
            configuration = request.form.get('configuration')
            
            if not integration_name:
                flash('Integration name is required', 'danger')
                return render_template('admin_ai_integration_create.html')
            
            try:
                AdminRepository.create_or_update_ai_integration(
                    integration_name, api_key, is_active, api_endpoint, configuration, current_user.id
                )
                flash('AI integration configured successfully!', 'success')
                return redirect(url_for('admin_ai_integrations'))
            except Exception as e:
                flash(f'Error configuring AI integration: {str(e)}', 'danger')
        
        return render_template('admin_ai_integration_create.html')
    
    @app.route('/admin/ai-integrations/<int:integration_id>/edit', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_edit_ai_integration(integration_id):
        integration = AdminRepository.get_ai_integration_by_id(integration_id)
        if not integration:
            flash('AI integration not found', 'danger')
            return redirect(url_for('admin_ai_integrations'))
        
        if request.method == 'POST':
            api_key = request.form.get('api_key')
            is_active = request.form.get('is_active') == 'on'
            api_endpoint = request.form.get('api_endpoint')
            configuration = request.form.get('configuration')
            
            # Only update API key if provided (not empty)
            api_key_to_update = api_key if api_key and api_key.strip() else None
            
            try:
                AdminRepository.create_or_update_ai_integration(
                    integration.integration_name, api_key_to_update, is_active, api_endpoint, configuration, current_user.id
                )
                flash('AI integration updated successfully!', 'success')
                return redirect(url_for('admin_ai_integrations'))
            except Exception as e:
                flash(f'Error updating AI integration: {str(e)}', 'danger')
        
        return render_template('admin_ai_integration_edit.html', integration=integration)
    
    @app.route('/admin/ai-integrations/<int:integration_id>/toggle', methods=['POST'])
    @role_required('Admin')
    def admin_toggle_ai_integration(integration_id):
        integration = AdminRepository.get_ai_integration_by_id(integration_id)
        if integration:
            try:
                AdminRepository.create_or_update_ai_integration(
                    integration.integration_name, None, not integration.is_active, 
                    None, None, current_user.id
                )
                status = "activated" if not integration.is_active else "deactivated"
                flash(f'AI integration {status} successfully!', 'success')
            except Exception as e:
                flash(f'Error toggling AI integration: {str(e)}', 'danger')
        
        return redirect(url_for('admin_ai_integrations'))
    
    @app.route('/api/admin/ai/toggle', methods=['POST'])
    @role_required('Admin')
    def admin_toggle_ai_enabled():
        """Toggle AI enabled/disabled state"""
        try:
            data = request.get_json() or {}
            enabled = data.get('enabled', True)
            
            # Get or create Gemini integration
            integration = AdminRepository.get_ai_integration_by_name('gemini')
            
            if integration:
                # Update existing integration
                integration.is_active = enabled
                integration.updated_by = current_user.id
                db.session.commit()
            else:
                # Create new integration record with enabled state
                AdminRepository.create_or_update_ai_integration(
                    'gemini', None, enabled, None, None, current_user.id
                )
            
            return jsonify({
                'success': True,
                'enabled': enabled,
                'message': f'AI features {"enabled" if enabled else "disabled"} successfully'
            })
        except Exception as e:
            return jsonify({
                'success': False,
                'message': f'Error toggling AI: {str(e)}'
            }), 500
    
    @app.route('/api/admin/ai/status', methods=['GET'])
    @role_required('Admin')
    def admin_ai_status():
        """Get Gemini AI integration status from environment and database"""
        try:
            from dotenv import load_dotenv
            import json
            
            load_dotenv()
            api_key = os.getenv('GEMINI_API_KEY')
            configured = bool(api_key)
            
            # Get database integration if exists
            db_integration = AdminRepository.get_ai_integration_by_name('gemini')
            
            enabled = False  # DB toggle flag (admin control)
            current_model = None
            
            if db_integration:
                enabled = db_integration.is_active
                if db_integration.configuration:
                    try:
                        config = json.loads(db_integration.configuration)
                        current_model = config.get('model')
                    except:
                        pass
            
            # Connected status: If configured (env var exists), Gemini is effectively connected and working
            # The backend uses the env var directly, so if it exists, the system is connected
            connected = configured
            
            return jsonify({
                'provider': 'gemini',
                'configured': configured,  # API key exists in env
                'enabled': enabled,      # DB toggle (admin preference)
                'connected': connected,  # Actually working (if configured, it's connected)
                'model': current_model
            })
                
        except Exception as e:
            return jsonify({
                'provider': 'gemini',
                'configured': False,
                'enabled': False,
                'connected': False,
                'model': None,
                'error': str(e)
            }), 500
    
    @app.route('/api/admin/ai/test', methods=['POST'])
    @role_required('Admin')
    def admin_test_ai_connection():
        """Test Gemini AI connection"""
        try:
            import google.generativeai as genai
            from dotenv import load_dotenv
            
            load_dotenv()
            api_key = os.getenv('GEMINI_API_KEY')
            
            if not api_key:
                return jsonify({
                    'success': False,
                    'message': 'GEMINI_API_KEY not found in environment variables'
                }), 400
            
            # Configure and test connection
            genai.configure(api_key=api_key)
            
            # Try to list models (lightweight test)
            try:
                models = genai.list_models()
                # Check if we can access at least one model
                model_found = False
                for model in models:
                    if hasattr(model, 'supported_generation_methods') and 'generateContent' in model.supported_generation_methods:
                        model_found = True
                        break
                
                if model_found:
                    return jsonify({
                        'success': True,
                        'message': 'Connection successful. Gemini API is accessible.'
                    })
                else:
                    return jsonify({
                        'success': False,
                        'message': 'No supported models found'
                    })
            except Exception as e:
                return jsonify({
                    'success': False,
                    'message': f'Failed to connect: {str(e)}'
                }), 500
                
        except Exception as e:
            return jsonify({
                'success': False,
                'message': f'Error testing connection: {str(e)}'
            }), 500
    
    # --- LMS Integration Routes (UC15, FR20) ---
    @app.route('/admin/lms-integrations')
    @role_required('Admin')
    def admin_lms_integrations():
        integrations = AdminRepository.get_all_lms_integrations()
        return render_template('admin_lms_integrations.html', integrations=integrations)
    
    @app.route('/admin/lms-integrations/create', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_create_lms_integration():
        if request.method == 'POST':
            lms_type = request.form.get('lms_type')
            lms_name = request.form.get('lms_name')
            api_url = request.form.get('api_url')
            api_key = request.form.get('api_key')
            api_secret = request.form.get('api_secret')
            course_id = request.form.get('course_id')
            is_active = request.form.get('is_active') == 'on'
            sync_enabled = request.form.get('sync_enabled') == 'on'
            configuration = request.form.get('configuration')
            
            if not lms_type or not lms_name or not api_url:
                flash('LMS type, name, and API URL are required', 'danger')
                return render_template('admin_lms_integration_create.html')
            
            try:
                AdminRepository.create_or_update_lms_integration(
                    lms_type, lms_name, api_url, api_key, api_secret, course_id,
                    is_active, sync_enabled, configuration, current_user.id
                )
                flash('LMS integration configured successfully!', 'success')
                return redirect(url_for('admin_lms_integrations'))
            except Exception as e:
                flash(f'Error configuring LMS integration: {str(e)}', 'danger')
        
        return render_template('admin_lms_integration_create.html')
    
    @app.route('/admin/lms-integrations/<int:integration_id>/edit', methods=['GET', 'POST'])
    @role_required('Admin')
    def admin_edit_lms_integration(integration_id):
        integration = AdminRepository.get_lms_integration_by_id(integration_id)
        if not integration:
            flash('LMS integration not found', 'danger')
            return redirect(url_for('admin_lms_integrations'))
        
        if request.method == 'POST':
            lms_name = request.form.get('lms_name')
            api_url = request.form.get('api_url')
            api_key = request.form.get('api_key')
            api_secret = request.form.get('api_secret')
            course_id = request.form.get('course_id')
            is_active = request.form.get('is_active') == 'on'
            sync_enabled = request.form.get('sync_enabled') == 'on'
            configuration = request.form.get('configuration')
            
            # Only update API key if provided (not empty)
            api_key_to_update = api_key if api_key and api_key.strip() else None
            api_secret_to_update = api_secret if api_secret and api_secret.strip() else None
            
            try:
                AdminRepository.create_or_update_lms_integration(
                    integration.lms_type, lms_name, api_url, api_key_to_update, api_secret_to_update,
                    course_id, is_active, sync_enabled, configuration, current_user.id
                )
                flash('LMS integration updated successfully!', 'success')
                return redirect(url_for('admin_lms_integrations'))
            except Exception as e:
                flash(f'Error updating LMS integration: {str(e)}', 'danger')
        
        return render_template('admin_lms_integration_edit.html', integration=integration)
    
    @app.route('/admin/lms-integrations/<int:integration_id>/delete', methods=['POST'])
    @role_required('Admin')
    def admin_delete_lms_integration(integration_id):
        try:
            AdminRepository.delete_lms_integration(integration_id)
            flash('LMS integration deleted successfully!', 'success')
        except Exception as e:
            flash(f'Error deleting LMS integration: {str(e)}', 'danger')
        
        return redirect(url_for('admin_lms_integrations'))
    
    @app.route('/admin/lms-integrations/<int:integration_id>/sync', methods=['POST'])
    @role_required('Admin')
    def admin_sync_lms_grades(integration_id):
        from services.lms_service import LMSService
        student_id = request.form.get('student_id', type=int)
        submission_id = request.form.get('submission_id', type=int)
        
        try:
            result = LMSService.sync_grades_to_lms(integration_id, student_id, submission_id)
            if result.get('success'):
                flash(result.get('message', 'Grades synced successfully!'), 'success')
            else:
                flash(result.get('message', 'Failed to sync grades'), 'danger')
        except Exception as e:
            flash(f'Error syncing grades: {str(e)}', 'danger')
        
        return redirect(url_for('admin_lms_integrations'))
    
    # --- Adaptive Insights Route (UC17) ---
    @app.route('/admin/generate-insights', methods=['POST'])
    @role_required('Admin')
    def admin_generate_insights():
        from services.adaptive_insights_service import AdaptiveInsightsService
        
        try:
            result = AdaptiveInsightsService.generate_insights_for_all_students()
            flash(f'Generated insights for {result["success_count"]}/{result["total_students"]} students. Total insights: {result["total_insights"]}', 'success')
        except Exception as e:
            flash(f'Error generating insights: {str(e)}', 'danger')
        
        return redirect(url_for('admin_dashboard'))

    return app

if __name__ == '__main__':
    app = create_app()
    app.run(debug=True)
